import os, json, re, base64, tempfile, shutil, copy, uuid, mimetypes, logging
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Optional, List
from fastapi import FastAPI, Query
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse, JSONResponse
from pydantic import BaseModel, Field, ConfigDict
import anthropic
import pdfplumber
import fitz  # pymupdf
import openpyxl
import zipfile
import docx as python_docx   # python-docx
import xlrd                   # legacy .xls support
from demand_service import compute_demand, market_position_score, build_market_audit, POSTCODE_AREA_KM2
from geospatial_competitors import get_nearby_competitors, material_supply_difference
from pipeline_supply import build_pipeline_supply
from run_diff import build_run_diff
from structured_deal import build_structured_deal_intelligence
from supabase import create_client

app = FastAPI()
logger = logging.getLogger("acquira-api")
logging.basicConfig(level=os.environ.get("LOG_LEVEL", "INFO"))

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # tighten to your Vercel URL in production
    allow_methods=["POST", "GET"],
    allow_headers=["*"],
)

MODEL      = "claude-sonnet-4-20250514"
MAX_TOKENS = 12000
API_RELEASE = "pipeline-retention-nonfatal-20260506"

client   = anthropic.Anthropic(api_key=os.environ["ANTHROPIC_API_KEY"])
supabase = create_client(os.environ["SUPABASE_URL"], os.environ["SUPABASE_SERVICE_KEY"])

# ── Prompts ───────────────────────────────────────────────────────────────────

EXTRACTION_SYSTEM_PROMPT = """You are an expert childcare acquisition analyst for Acquira, an Australian deal intelligence platform. Your task is to read an Information Memorandum (IM) for a childcare centre and extract structured data into a strict JSON format.

RULES:
1. Return ONLY valid JSON. No preamble, no explanation, no markdown code fences.
2. Never invent or estimate numbers. If a field is not present in the IM, set it to null.
3. Use null (not zero, not "N/A", not "unknown") for any field you cannot find.
4. Be conservative: if you are uncertain whether a number is correct, set it to null and add the field name to the missing_fields array.
5. All dollar amounts in AUD as plain numbers (e.g. 1250000 not "$1.25M"). All percentages as plain numbers 0-100 (e.g. 78.5 not 0.785). All dates as ISO 8601 strings (YYYY-MM-DD).
6. You are state-aware: VIC, NSW, and QLD have different kinder/preschool funding regimes, regulatory bodies, and wage award rates.
7. For occupancy: use the most recent data available. Priority order: (1) most recent 4-week average, (2) most recent month, (3) annual average.
8. For financials: extract ALL years available (FY23, FY24, FY25). Prefer audited or management accounts over vendor summaries.
9. RATIO CALCULATIONS - always calculate from raw numbers:
   - labour_ratio_pct = (total_labour_cost / revenue) x 100
   - rent_ratio_pct = (rent_pa / revenue) x 100
   - ebitda_margin_pct = (ebitda / revenue) x 100
   - ebitda = revenue - total_labour_cost - rent_pa - other_operating_costs
10. For asking price: if the IM states "Price on Application", "POA" -> set asking_price to null and add "asking_price_poa": true to meta.
11. Flag unusual patterns or vendor-inflated items in the anomalies array.
12. Return all numbers as plain JSON numbers. Never use a leading + sign.
13. labour_ratio_pct >100 or <20 is almost certainly an error - recheck and set to null if unresolvable.
14. GREENFIELD / PRE-OPENING RULE: If the document describes a brand-new, turn-key, approaching-OC, not-yet-operating, AFL assignment, or lease assignment centre, classify it as greenfield_pre_opening in meta.source_type. Do NOT treat pro-forma EBITDA as actual FY25 EBITDA. Historical occupancy, labour ratio, NQS, and actual EBITDA should remain null unless explicitly provided. Put vendor forecast/pro-forma EBITDA values in meta.anomalies and/or pipeline_mentions rather than fy25.ebitda.

VIC-SPECIFIC CONTEXT:
- VKF (Vic Kinder Funding) = State Government kindergarten subsidy, counted as revenue.
- NQS ratings: Excellent > Exceeding NQS > Meeting NQS > Working Towards NQS > Significant Improvement Required.
- Owner-operator director wages are a standard addback item.
- ADDBACK RULES: Normalised EBITDA = reported EBITDA + verified addbacks.
  Standard addbacks to extract and itemise:
  (1) Owner/director salary above market replacement cost (~$80-110K for a centre manager)
  (2) One-off non-recurring expenses clearly stated in the IM
  (3) Personal expenses run through the business (must be stated in IM)
  Do NOT invent addbacks. Only include what is explicitly stated in the documents.
- Always extract BOTH: reported_ebitda (as stated) AND normalised_ebitda (after addbacks).
- If addbacks are claimed but unverified, flag them in anomalies with reason.
- Labour ratio and rent ratio should be calculated on REPORTED revenue, not normalised.

Return this exact JSON structure (set fields to null if not found):

{
  "meta": {
    "extraction_version": "1.1",
    "extraction_date": "",
    "source_type": "pdf_im",
    "source_files": [],
    "data_quality": "MEDIUM",
    "missing_fields_count": 0,
    "missing_fields": [],
    "asking_price_poa": false,
    "anomalies": []
  },
  "centre": {
    "name": null, "trading_name": null, "address": null, "suburb": null,
    "state": null, "postcode": null, "lga": null, "operator": null,
    "operator_type": "unknown", "licensed_places": null, "nqs_rating": null,
    "nqs_date": null, "service_approval_number": null
  },
  "occupancy": {
    "current_month_pct": null, "avg_4wk_pct": null, "avg_13wk_pct": null,
    "avg_52wk_pct": null, "peak_pct": null, "peak_week": null,
    "fy23_avg_pct": null, "fy24_avg_pct": null, "fy25_avg_pct": null,
    "trend_fy23_to_fy25": null, "waitlist_depth": null, "waitlist_notes": null
  },
  "financials": {
    "primary_year": "FY25",
    "fy23": {"revenue": null, "total_labour_cost": null, "rent_pa": null, "ebitda": null, "labour_ratio_pct": null, "rent_ratio_pct": null, "ebitda_margin_pct": null},
    "fy24": {"revenue": null, "total_labour_cost": null, "rent_pa": null, "ebitda": null, "labour_ratio_pct": null, "rent_ratio_pct": null, "ebitda_margin_pct": null},
    "fy25": {"revenue": null, "total_labour_cost": null, "rent_pa": null, "ebitda": null, "labour_ratio_pct": null, "rent_ratio_pct": null, "ebitda_margin_pct": null},
    "ebitda_3yr_average": null, "revenue_trend": null, "labour_trend": null,
    "asking_price": null, "asking_price_ebitda_multiple": null,
    "addbacks": {
      "owner_salary_addback": null,
      "owner_salary_note": null,
      "other_addbacks": [],
      "addbacks_total": null,
      "reported_ebitda": null,
      "normalised_ebitda": null,
      "normalised_ebitda_multiple": null,
      "addback_confidence": "none"
    }
  },
  "lease": {
    "commencement_date": null, "expiry_date": null, "status": "UNKNOWN",
    "term_years": null, "options": null, "remaining_term_years": null,
    "base_rent_pa_fy25": null, "rent_review_type": null, "rent_review_detail": null,
    "turnover_rent_clause": null, "assignment_clause": null,
    "demolition_redevelopment_clause": null, "make_good_obligations": null,
    "outgoings_type": null, "permitted_use": null, "lessor": null, "lessee": null
  },
  "hard_flags": [],
  "key_ratios": {
    "occupancy_latest_4wk_pct": null, "occupancy_peak_pct": null,
    "revenue_fy25": null, "ebitda_fy25": null, "ebitda_margin_fy25_pct": null,
    "labour_ratio_fy25_pct": null, "rent_ratio_fy25_pct": null,
    "ebitda_3yr_avg": null, "rent_pa_fy25": null, "licensed_places": null,
    "asking_price": null, "ebitda_multiple": null
  },
  "pipeline_mentions": []
}

PIPELINE MENTIONS: Extract any mentions of:
- Approved DAs or council planning applications for childcare/early learning nearby
- Competitor centres under construction or recently opened
- Sites approved for childcare development in the area
- Any new supply risks mentioned in the document
Store as an array of strings, e.g. ["DA approved at 45 Smith St for 90-place centre", "New childcare under construction 500m away"]
If none found, return an empty array []."""

SCORING_SYSTEM_PROMPT = """You are an expert childcare acquisition analyst for Acquira. You receive structured data extracted from a childcare centre Information Memorandum (IM) and score it across 17 dimensions.

ABSOLUTE OUTPUT RULES:
1. Return ONLY valid JSON. No preamble, no markdown fences, no text outside the JSON.
2. All numbers must be plain JSON numbers. Never use a leading + sign.
3. Use null for any value that cannot be determined — never omit a key.
4. temperature is 0 — your output must be fully deterministic given the same input.

SCORING PHILOSOPHY:
- Score each dimension on a 0-10 scale using the rubric below.
- 5.0 = industry average / neutral. 7.0+ = genuinely good. 9.0+ = exceptional.
- Every score MUST quote the actual number from the data in its summary.
- Dimension summaries must be 2-3 sentences specific to THIS deal.
- The server will recalculate total_score using weights — set it to 0 in your output.
- If data is missing for a dimension, score it 5.0 (neutral) and note it in summary.

DIMENSION WEIGHTS (for reference — server recalculates):
occupancy_demand:        0.15
profitability_cashflow:  0.15
revenue_pricing:         0.08
staffing_resilience:     0.08
lease_economics:         0.08
valuation_structure:     0.08
market_position:         0.07
management_systems:      0.04
regulatory_quality:      0.05
upside_levers:           0.03
ccs_risk:                0.07
lease_tail:              0.03
capex_liability:         0.02
staff_qualification_mix: 0.02
fee_benchmarking:        0.02
operator_quality:        0.02
enrolment_trend:         0.01

SCORING RUBRIC (0-10 per dimension):

occupancy_demand:
  9-10: occ >= 90%, strong waitlist
  7-8:  occ 75-89%, some waitlist
  5-6:  occ 60-74%, stable
  3-4:  occ 45-59%, declining or flat
  1-2:  occ < 45%, critical

profitability_cashflow:
  9-10: EBITDA margin >= 25%, positive 3yr trend
  7-8:  margin 18-24%
  5-6:  margin 10-17%
  3-4:  margin 3-9%
  1-2:  margin < 3% or negative
  IMPORTANT: If 3yr average EBITDA is negative even when FY25 is positive,
  cap this score at 4.0 maximum and note the sustained loss history.
  Use normalised_ebitda if addbacks are verified.
  Stress test: model a +10% wage cost increase (apply to total_labour_cost).
  If the resulting EBITDA drops below 0, flag HIGH RISK: Wage stress wipes EBITDA.
  If the resulting defensive yield (stressed EBITDA / asking_price) falls below 4.5%, flag HIGH RISK: Defensive yield below 4.5% under wage stress.
  State the stressed EBITDA and stressed yield in your summary.

staffing_resilience:
  9-10: labour ratio < 52%
  7-8:  labour 52-58%
  5-6:  labour 58-65%
  3-4:  labour 65-72%
  1-2:  labour > 72%

lease_economics:
  9-10: rent < 10% of revenue
  7-8:  rent 10-15%
  5-6:  rent 15-20%
  3-4:  rent 20-25%
  1-2:  rent > 25%

lease_tail:
  9-10: >= 15 years remaining tenure (incl options)
  7-8:  10-14 years
  5-6:  5-9 years
  3-4:  2-4 years
  1-2:  < 2 years or expired

regulatory_quality:
  9-10: Exceeding NQS, recent assessment
  7-8:  Meeting NQS
  5-6:  Meeting NQS, assessment overdue
  3-4:  Working Towards NQS
  1-2:  SIR or active compliance notices

valuation_structure:
  9-10: <= 2x EBITDA
  7-8:  2-3x EBITDA
  5-6:  3-4x EBITDA
  3-4:  4-5x EBITDA
  1-2:  > 5x EBITDA or POA

market_position:
  9-10: < 1.5 competitors per licensed place within 3km, strong demographics
  7-8:  balanced supply zone
  5-6:  average competition
  3-4:  oversupplied market
  1-2:  heavily oversupplied
  Saturation penalty: apply a haircut to your initial score based on competitor count within 3km:
    1-2 competitors: -10% (multiply score by 0.90)
    3 competitors:   -15% (multiply score by 0.85)
    4-5 competitors: -25% (multiply score by 0.75)
    6+ competitors:  -35% (multiply score by 0.65)
  State the multiplier used in your summary.
  Approved-but-unbuilt centres: if the IM or context mentions DA-approved or under-construction centres nearby,
  apply an additional -20% penalty and flag Pipeline supply risk.
  DA Pipeline: if the user has provided pipeline_intel (approved DAs, lodged applications, permit sites for sale),
  incorporate this into the market_position score.
  Approved DAs within 3km that add >25% of the centre's licensed places = HIGH pipeline risk, reduce score by 2 points minimum.
  Approved DAs >50% of licensed places = CRITICAL pipeline risk, reduce by 3-4 points.
  State the pipeline risk explicitly in the summary (e.g. "HIGH pipeline risk: X approved DAs add Y places within 3km").

management_systems:
  7-8:  professional management team, strong systems, not owner-dependent
  5-6:  semi-professional, some key-person risk
  3-4:  owner-operator, high transition risk
  1-2:  sole operator, no documented systems

revenue_pricing:
  7-8:  fees above suburb median, strong CCS mix, multiple revenue streams
  5-6:  fees at market, standard CCS dependency
  3-4:  fees below market, limited pricing power
  1-2:  fees significantly below market, no uplift path

upside_levers:
  7-8:  clear fee uplift headroom, occupancy expansion possible, kinder funding upside
  5-6:  some upside, limited by market or occupancy ceiling
  3-4:  limited upside, near capacity or market ceiling
  1-2:  no meaningful upside identified

ccs_risk:
  7-8:  low CCS cliff exposure, diverse family demographics
  5-6:  moderate CCS dependency
  3-4:  high CCS dependency, activity test risk
  1-2:  critical CCS exposure
  Cohort risk: estimate the % of revenue derived from the 3-5 age group (kindy/preschool cohort).
  If that % is >40% AND the IM mentions Pre-Prep, kindy, or preschool competition nearby,
  flag HIGH RISK: Systematic cohort loss risk. Add cohort_35_pct_estimated to the detail block.
  Revenue topology: classify the CCS revenue routing as one of:
    Direct (government pays centre directly) → LOW risk,
    Parent-routed (subsidy flows via parent) → MEDIUM risk,
    Program-dependent (tied to a specific govt program that could end) → HIGH risk.
  State the classification and rationale in your summary.

capex_liability:
  9-10: new fit-out, no CAPEX required
  7-8:  < 3 years old, minimal CAPEX
  5-6:  3-7 years, routine maintenance only
  3-4:  7-12 years, significant refresh likely
  1-2:  > 12 years or major CAPEX flagged

staff_qualification_mix:
  7-8:  > 35% degree qualified, low wage trajectory risk
  5-6:  20-35% degree qualified
  3-4:  < 20% degree qualified, high wage risk
  1-2:  unknown or significant compliance risk

fee_benchmarking:
  7-8:  fees >= suburb median, room to increase
  5-6:  fees within 5% of median
  3-4:  fees > 5% below median
  1-2:  fees significantly below market
  Fee ceiling risk: if the centre's daily fee is at or above the suburb's top quartile (estimated as median + 10%)
  AND no premium differentiation is noted (e.g. no Reggio, Forest School, specialty program, Exceeding NQS),
  flag Fee ceiling risk and cap this dimension score at 5.0 maximum.

operator_quality:
  9-10: Exceeding NQS, no conditions/notices, strong compliance history
  7-8:  Meeting NQS, clean record
  5-6:  Meeting NQS, minor issues resolved
  3-4:  Working Towards NQS or unresolved conditions
  1-2:  SIR, active notices, enforcement history

enrolment_trend:
  9-10: occupancy improving, strong waitlist across age groups
  7-8:  stable at high occupancy, moderate waitlist
  5-6:  stable, no waitlist
  3-4:  declining occupancy trend
  1-2:  significant decline

DEAL-BREAKER FLAGS to evaluate (set triggered true/false):
occupancy_critical (occ<50%), occupancy_warning (50-65%), rent_ratio_danger (rent>15% rev),
labour_ratio_danger (labour>65% rev), ebitda_negative, lease_short_no_options (<3yr, no options),
lease_short_with_options (<3yr with options), owner_operator_dependency,
nqs_working_towards, capex_high, ccs_exposure_high, valuation_premium (>4x EBITDA turnaround)

Return this exact JSON schema (null for unknowns, never omit keys):
{
  "centre_name": string,
  "total_score": number,
  "dimensions": {
    "occupancy_demand":       {"score": 0-10, "label": string, "summary": string, "data_used": []},
    "revenue_pricing":        {"score": 0-10, "label": string, "summary": string, "data_used": []},
    "staffing_resilience":    {"score": 0-10, "label": string, "summary": string, "data_used": []},
    "profitability_cashflow": {"score": 0-10, "label": string, "summary": string, "data_used": []},
    "lease_economics":        {"score": 0-10, "label": string, "summary": string, "data_used": []},
    "regulatory_quality":     {"score": 0-10, "label": string, "summary": string, "data_used": []},
    "market_position":        {"score": 0-10, "label": string, "summary": string, "data_used": []},
    "management_systems":     {"score": 0-10, "label": string, "summary": string, "data_used": []},
    "valuation_structure":    {"score": 0-10, "label": string, "summary": string, "data_used": []},
    "upside_levers":          {"score": 0-10, "label": string, "summary": string, "data_used": []},
    "ccs_risk":               {"score": 0-10, "label": "CCS / Subsidy Risk", "summary": string, "data_used": [], "detail": {"estimated_ccs_dependent_pct": null, "activity_test_exposure": "unknown", "subsidy_cliff_note": string}},
    "lease_tail":             {"score": 0-10, "label": "Lease Tail", "summary": string, "data_used": [], "detail": {"years_remaining": null, "options_available": null, "option_years_each": null, "total_potential_tenure": null, "landlord_obligations_noted": null}},
    "capex_liability":        {"score": 0-10, "label": "Renovation / CAPEX Liability", "summary": string, "data_used": [], "detail": {"fit_out_age_years": null, "capex_mentioned_in_im": false, "estimated_capex_risk": "unknown", "notes": string}},
    "staff_qualification_mix":{"score": 0-10, "label": "Staff Qualification Mix", "summary": string, "data_used": [], "detail": {"degree_qualified_pct": null, "certificate_pct": null, "diploma_pct": null, "wage_trajectory_risk": "unknown"}},
    "fee_benchmarking":       {"score": 0-10, "label": "Fee Benchmarking", "summary": string, "data_used": [], "detail": {"centre_daily_fee": null, "suburb_median_fee": null, "fee_position": "unknown", "pricing_power_note": string}},
    "operator_quality":       {"score": 0-10, "label": "Operator Quality Signal", "summary": string, "data_used": [], "detail": {"nqs_rating": "unknown", "last_assessment_date": null, "months_since_assessment": null, "exceeding_areas_count": null, "active_conditions": null, "active_notices": null, "compliance_note": string}},
    "enrolment_trend":        {"score": 0-10, "label": "Enrolment Trend & Waitlist", "summary": string, "data_used": [], "detail": {"current_occupancy_pct": null, "trend_direction": "unknown", "waitlist_depth": "unknown", "occupancy_snapshot_date": null, "trend_note": string}}
  },
  "deal_breaker_flags": {
    "any_triggered": false,
    "flags": [{"id": string, "triggered": bool, "severity": "critical"|"high", "label": string, "reason": string}]
  },
  "audit_trail": {
    "fields_missing": [],
    "confidence": "medium",
    "confidence_note": string
  },
  "verdict": {
    "category": "passive_hold"|"turnaround"|"distressed"|"pass",
    "one_liner": string,
    "recommended_buyer_profile": string
  }
}"""

# ── Helpers ───────────────────────────────────────────────────────────────────

def classify_file(filename: str) -> str:
    f = filename.lower()
    # Check extension first to avoid e.g. 'lease terms.docx' -> lease_pdf
    if f.endswith(('.xlsx', '.xls', '.csv')):
        if any(x in f for x in ['p&l', 'p_l', 'profit', 'loss']): return 'pl_excel'
        if any(x in f for x in ['occupancy', 'utilisation', 'utilization']): return 'occupancy_excel'
        if 'transaction' in f: return 'transaction_excel'
        if 'payroll' in f: return 'payroll_excel'
        return 'pl_excel'
    if f.endswith('.pdf'):
        if any(x in f for x in ['lease', 'deed', 'tenancy']): return 'lease_pdf'
        if 'service approval' in f: return 'service_approval_pdf'
        if any(x in f for x in ['nqs', 'acecqa', 'rating']): return 'nqs_pdf'
        return 'im_pdf'
    if f.endswith('.docx'):
        if any(x in f for x in ['lease', 'deed', 'tenancy']): return 'lease_docx'
        return 'im_docx'
    return 'unknown'

def extract_pdf_text(pdf_path: str) -> str:
    try:
        text = ''
        with pdfplumber.open(pdf_path) as pdf:
            for idx, page in enumerate(pdf.pages, start=1):
                t = page.extract_text()
                if t:
                    text += f'\n--- Page {idx} ---\n{t}\n'
        return text[:80000]
    except Exception:
        return ''

def is_pdf_scanned(text: str) -> bool:
    trimmed = text.strip()
    if len(trimmed) < 200:
        return True
    has_dollars = bool(re.search(r'\$[\d,]+|[\d,]+\s*(revenue|ebitda|wages|labour|rent)', trimmed, re.I))
    avg_chars = len(trimmed) / max(len(trimmed.split('\n\n')), 1)
    return avg_chars < 300 and not has_dollars

async def extract_scanned_pdf_text(pdf_path: str, purpose: str) -> str:
    try:
        doc = fitz.open(pdf_path)
        images = []
        for i, page in enumerate(doc):
            if i >= 60: break
            if len(page.get_text().strip()) < 30 and i > 3: continue
            mat = fitz.Matrix(1.5, 1.5)
            pix = page.get_pixmap(matrix=mat)
            images.append(base64.standard_b64encode(pix.tobytes('png')).decode())
            if len(images) >= 30: break

        if not images:
            return ''

        content = [
            *[{"type": "image", "source": {"type": "base64", "media_type": "image/png", "data": img}} for img in images],
            {"type": "text", "text": f"Extract all text content from these document pages. This is a {purpose}. Return plain text only, preserving structure and numbers accurately. Prefix each page in order with a marker like --- Page 1 ---, --- Page 2 ---, etc."}
        ]

        response = client.messages.create(
            model=MODEL,
            max_tokens=8000,
            temperature=0,
            messages=[{"role": "user", "content": content}]
        )
        return response.content[0].text if response.content[0].type == 'text' else ''
    except Exception as e:
        print(f"Vision extraction failed: {e}")
        return ''

def extract_excel_text(xlsx_path: str) -> str:
    """Extract text from .xlsx or legacy .xls files."""
    path_lower = xlsx_path.lower()
    try:
        if path_lower.endswith('.xls') and not path_lower.endswith('.xlsx'):
            # Legacy BIFF format — openpyxl cannot read this
            wb = xlrd.open_workbook(xlsx_path)
            out = []
            for sheet in wb.sheets():
                out.append(f'Sheet: {sheet.name}')
                for rx in range(sheet.nrows):
                    row = [str(sheet.cell_value(rx, cx)) for cx in range(sheet.ncols)]
                    line = ','.join(v for v in row if v.strip())
                    if line:
                        out.append(line)
                    if len(out) >= 1000:
                        break
            return '\n'.join(out[:1000])
        else:
            wb = openpyxl.load_workbook(xlsx_path, data_only=True)
            out = []
            for sheet in wb.sheetnames:
                ws = wb[sheet]
                out.append(f'Sheet: {sheet}')
                for row in ws.iter_rows(values_only=True):
                    if any(v is not None for v in row):
                        out.append(','.join(str(v) if v is not None else '' for v in row))
                    if len(out) >= 1000:
                        break
            return '\n'.join(out[:1000])
    except Exception as e:
        print(f"Excel extraction failed ({xlsx_path}): {e}")
        return ''

def extract_docx_text(docx_path: str) -> str:
    """Extract text from .docx using the public python-docx API."""
    try:
        doc = python_docx.Document(docx_path)
        parts = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                parts.append(text)
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                seen = []
                for c in cells:
                    if c and (not seen or c != seen[-1]):
                        seen.append(c)
                line = '\t'.join(seen)
                if line:
                    parts.append(line)
        return '\n'.join(parts)
    except Exception as e:
        print(f"DOCX extraction failed: {e}")
        return ''

def _extract_file_text(file_path: str, filename: str) -> str:
    """Route a file to the correct extractor by extension."""
    f = filename.lower()
    if f.endswith('.pdf'):
        return extract_pdf_text(file_path)
    elif f.endswith('.docx'):
        return extract_docx_text(file_path)
    elif f.endswith(('.xlsx', '.xls', '.csv')):
        return extract_excel_text(file_path)
    return ''

def clean_json(text: str) -> str:
    text = re.sub(r'^```json\s*', '', text, flags=re.M)
    text = re.sub(r'^```\s*', '', text, flags=re.M)
    text = re.sub(r'```$', '', text, flags=re.M)
    text = re.sub(r':\s*\+([0-9])', r': \1', text)
    return text.strip()

# ── SSE helper ────────────────────────────────────────────────────────────────

def sse_event(event: str, data: dict) -> str:
    return f"event: {event}\ndata: {json.dumps(data)}\n\n"

def normalize_storage_path(path: str) -> str:
    """
    Supabase storage.download() requires the object path INSIDE the bucket,
    not a public/signed URL and not a path prefixed with the bucket name.

    Examples accepted by this helper:
      - user-id/file.pdf
      - uploads/user-id/file.pdf
      - https://.../storage/v1/object/public/uploads/user-id/file.pdf
      - https://.../storage/v1/object/sign/uploads/user-id/file.pdf?token=...
    """
    if not path:
        return path

    path = str(path).strip()

    markers = [
        "/storage/v1/object/public/uploads/",
        "/storage/v1/object/sign/uploads/",
        "/storage/v1/object/authenticated/uploads/",
    ]
    for marker in markers:
        if marker in path:
            return path.split(marker, 1)[1].split("?", 1)[0]

    if path.startswith("uploads/"):
        return path[len("uploads/"):]

    return path

# ── Request model ─────────────────────────────────────────────────────────────

class PipelineIntel(BaseModel):
    # Accept both frontend camelCase and backend snake_case payloads.
    model_config = ConfigDict(populate_by_name=True)

    approved_das: Optional[int] = Field(default=None, alias="approvedDAs")
    lodged_applications: Optional[int] = Field(default=None, alias="lodgedApplications")
    permit_sites: Optional[int] = Field(default=None, alias="permitSites")
    notes: Optional[str] = None

class PipelineProject(BaseModel):
    id: Optional[str] = None
    name: Optional[str] = None
    address: Optional[str] = None
    distance_km: Optional[float] = None
    status: str = "unknown"
    proposed_places: Optional[Any] = None
    source_url: Optional[str] = None
    source_file: Optional[str] = None
    source_date: Optional[str] = None
    confidence: Optional[str] = None
    notes: Optional[str] = None

class PipelineRequest(BaseModel):
    model_config = ConfigDict(populate_by_name=True)

    # Multi-file (new)
    storagePaths: Optional[list[str]] = None
    filenames:    Optional[list[str]] = None
    # Single-file (legacy — backwards compat)
    storagePath:  Optional[str]       = None
    filename:     Optional[str]       = None
    # Pipeline intelligence (optional)
    pipelineIntel: Optional[PipelineIntel] = None
    pipelineProjects: Optional[list[PipelineProject]] = None

    def resolved_paths(self) -> list[str]:
        if self.storagePaths:
            return [normalize_storage_path(p) for p in self.storagePaths]
        if self.storagePath:
            return [normalize_storage_path(self.storagePath)]
        raise ValueError("No storage path provided")

    def resolved_filenames(self) -> list[str]:
        if self.filenames:
            return self.filenames
        if self.filename:
            return [self.filename]
        return [p.split('/')[-1] for p in self.resolved_paths()]

class ReunderwriteBase(BaseModel):
    extracted: dict[str, Any]
    scored: dict[str, Any]
    workflow: dict[str, Any]

class SelectedDiligenceDocument(BaseModel):
    id: str
    storage_path: str
    filename: str
    document_type: Optional[str] = None
    source_item_id: Optional[str] = None
    file_size: Optional[int] = None

class SelectedSourceDocument(BaseModel):
    id: str
    storage_path: str
    filename: str
    content_type: Optional[str] = None
    source_kind: Optional[str] = None
    run_id: Optional[str] = None
    file_size: Optional[int] = None

class ReunderwriteRequest(BaseModel):
    deal_id: str
    run_id: str
    base_run_id: str
    base: ReunderwriteBase
    selected_diligence_documents: Optional[list[SelectedDiligenceDocument]] = None
    selected_source_documents: Optional[list[SelectedSourceDocument]] = None
    input_document_count: Optional[int] = None
    input_total_bytes: Optional[int] = None
    pipeline_projects: Optional[list[PipelineProject]] = None
    mode: str = "reunderwrite"

# ── Routes ────────────────────────────────────────────────────────────────────

@app.get("/health")
def health():
    return {"status": "ok", "release": API_RELEASE}

REUNDERWRITE_TEXT_BUDGET: dict[str, tuple[int, int]] = {
    'im_pdf':               (20000, 3),
    'im_docx':              (20000, 3),
    'pl_excel':             (12000, 10),
    'occupancy_excel':      (12000, 10),
    'transaction_excel':    (8000,  5),
    'payroll_excel':        (12000, 10),
    'lease_pdf':            (10000, 5),
    'lease_docx':           (10000, 5),
    'service_approval_pdf': (5000,  3),
    'nqs_pdf':              (5000,  3),
}
CLAUDE_CHAR_LIMIT = 120_000
MAX_REUNDERWRITE_DOCUMENTS = 10
MAX_REUNDERWRITE_DECLARED_BYTES = 75 * 1024 * 1024
DIMENSION_WEIGHTS = {
    'occupancy_demand':        0.15,
    'profitability_cashflow':  0.15,
    'revenue_pricing':         0.08,
    'staffing_resilience':     0.08,
    'lease_economics':         0.08,
    'valuation_structure':     0.08,
    'market_position':         0.10,
    'management_systems':      0.04,
    'regulatory_quality':      0.05,
    'upside_levers':           0.00,
    'ccs_risk':                0.07,
    'lease_tail':              0.03,
    'capex_liability':         0.02,
    'staff_qualification_mix': 0.02,
    'fee_benchmarking':        0.02,
    'operator_quality':        0.02,
    'enrolment_trend':         0.01,
}

def _is_present(value: Any) -> bool:
    return value is not None and value != "" and value != [] and value != {}

def validate_diligence_storage_path(storage_path: str, deal_id: str) -> None:
    prefix = f"diligence/{deal_id}/"
    if not storage_path.startswith(prefix):
        raise ValueError("storage_path must be under diligence/{deal_id}/")
    remainder = storage_path[len(prefix):]
    if not remainder:
        raise ValueError("storage_path filename segment is required")
    if ".." in remainder or "\\" in remainder:
        raise ValueError("storage_path contains an unsafe segment")
    if any(ord(ch) < 32 or ord(ch) == 127 for ch in remainder):
        raise ValueError("storage_path contains control characters")
    if any(part == "" for part in remainder.split("/")):
        raise ValueError("storage_path contains empty path segments")

def validate_source_storage_path(storage_path: str) -> None:
    prefix = "deal-sources/"
    if not storage_path.startswith(prefix):
        raise ValueError("selected source storage_path must be under deal-sources/")
    remainder = storage_path[len(prefix):]
    if not remainder:
        raise ValueError("selected source storage_path requires a filename segment")
    if ".." in remainder or "\\" in remainder:
        raise ValueError("selected source storage_path contains an unsafe segment")
    if any(ord(ch) < 32 or ord(ch) == 127 for ch in remainder):
        raise ValueError("selected source storage_path contains control characters")
    if any(part == "" for part in remainder.split("/")):
        raise ValueError("selected source storage_path contains empty path segments")

def _safe_local_name(name: str) -> str:
    return re.sub(r'[^a-zA-Z0-9._-]', '_', Path(name).name or "diligence_document")

def _safe_storage_name(name: str, fallback: str = "source_document") -> str:
    safe = re.sub(r'[^a-zA-Z0-9._-]', '_', Path(name).name or fallback).strip("._")
    return safe[:180] or fallback

def validate_pipeline_temp_path(storage_path: str) -> None:
    if not storage_path.startswith("pipeline/"):
        raise ValueError("Initial pipeline source storage paths must be under pipeline/")
    remainder = storage_path[len("pipeline/"):]
    if not remainder:
        raise ValueError("Pipeline source storage path filename segment is required")
    if ".." in remainder or "\\" in remainder:
        raise ValueError("Pipeline source storage path contains an unsafe segment")
    if any(ord(ch) < 32 or ord(ch) == 127 for ch in remainder):
        raise ValueError("Pipeline source storage path contains control characters")
    if any(part == "" for part in remainder.split("/")):
        raise ValueError("Pipeline source storage path contains empty path segments")

def retain_pipeline_source_file(
    *,
    pipeline_request_id: str,
    index: int,
    original_storage_path: str,
    filename: str,
    file_bytes: bytes,
) -> dict[str, Any]:
    validate_pipeline_temp_path(original_storage_path)
    safe_filename = _safe_storage_name(filename, f"source_{index + 1}")
    retained_storage_path = f"deal-sources/pending/{pipeline_request_id}/{index + 1}-{safe_filename}"
    content_type = mimetypes.guess_type(filename)[0] or "application/octet-stream"
    supabase.storage.from_("uploads").upload(
        retained_storage_path,
        file_bytes,
        {"content-type": content_type, "x-upsert": "true"},
    )
    return {
        "original_storage_path": original_storage_path,
        "retained_storage_path": retained_storage_path,
        "filename": filename,
        "content_type": content_type,
        "file_size": len(file_bytes),
    }

async def parse_diligence_documents(
    docs: list[SelectedDiligenceDocument],
    *,
    deal_id: str,
    work_dir: str,
) -> tuple[str, list[str], dict[str, str], list[str]]:
    combined_text = ""
    source_files: list[str] = []
    file_classes: dict[str, str] = {}
    skipped: list[str] = []
    class_counts: dict[str, int] = {}

    for index, doc in enumerate(docs):
        validate_diligence_storage_path(doc.storage_path, deal_id)
        if not doc.id or not doc.filename:
            raise ValueError("Each selected diligence document requires id and filename")
        file_bytes = supabase.storage.from_('uploads').download(doc.storage_path)
        filename = doc.filename
        fname_lower = filename.lower()
        doc_label = f"{filename} [diligence:{doc.id}]"

        def add_text(base_name: str, file_class: str, text: str, label: str) -> None:
            nonlocal combined_text
            if not text:
                skipped.append(label)
                return
            max_chars, _max_count = REUNDERWRITE_TEXT_BUDGET.get(file_class, (6000, 3))
            source_files.append(label)
            file_classes[label] = file_class
            combined_text += (
                f"\n\n=== {label} ({file_class}) ===\n"
                "DOCUMENT_ROLE: Diligence document\n"
                f"DILIGENCE_DOCUMENT_ID: {doc.id}\n"
                f"STORAGE_PATH: {doc.storage_path}\n"
                f"SOURCE_ITEM_ID: {doc.source_item_id or 'none'}\n"
                f"DOCUMENT_TYPE: {doc.document_type or 'unknown'}\n\n"
                f"{text[:max_chars]}"
            )

        if fname_lower.endswith(".zip"):
            zip_path = os.path.join(work_dir, f"{index}_diligence.zip")
            with open(zip_path, "wb") as f:
                f.write(file_bytes)
            with zipfile.ZipFile(zip_path, "r") as zf:
                for entry_name in zf.namelist():
                    base_name = Path(entry_name).name
                    if not base_name or base_name.startswith("."):
                        continue
                    file_class = classify_file(base_name)
                    if file_class == "unknown":
                        skipped.append(base_name)
                        continue
                    _max_chars, max_count = REUNDERWRITE_TEXT_BUDGET.get(file_class, (6000, 3))
                    if class_counts.get(file_class, 0) >= max_count:
                        skipped.append(base_name)
                        continue
                    entry_path = os.path.join(work_dir, f"{index}_{_safe_local_name(base_name)}")
                    with open(entry_path, "wb") as f:
                        f.write(zf.read(entry_name))
                    text = _extract_file_text(entry_path, base_name)
                    if base_name.lower().endswith(".pdf") and is_pdf_scanned(text):
                        text = await extract_scanned_pdf_text(entry_path, file_class.replace("_", " "))
                    class_counts[file_class] = class_counts.get(file_class, 0) + 1
                    add_text(base_name, file_class, text, f"{base_name} [diligence:{doc.id}]")
            continue

        file_class = classify_file(filename)
        if file_class == "unknown":
            skipped.append(filename)
            continue
        _max_chars, max_count = REUNDERWRITE_TEXT_BUDGET.get(file_class, (10000, 3))
        if class_counts.get(file_class, 0) >= max_count:
            skipped.append(filename)
            continue
        file_path = os.path.join(work_dir, f"{index}_{_safe_local_name(filename)}")
        with open(file_path, "wb") as f:
            f.write(file_bytes)
        text = _extract_file_text(file_path, filename)
        if fname_lower.endswith(".pdf") and is_pdf_scanned(text):
            text = await extract_scanned_pdf_text(file_path, file_class.replace("_", " "))
        class_counts[file_class] = class_counts.get(file_class, 0) + 1
        add_text(filename, file_class, text, doc_label)

    return combined_text, source_files, file_classes, skipped

async def parse_source_documents(
    docs: list[SelectedSourceDocument],
    *,
    work_dir: str,
) -> tuple[str, list[str], dict[str, str], list[str]]:
    combined_text = ""
    source_files: list[str] = []
    file_classes: dict[str, str] = {}
    skipped: list[str] = []
    class_counts: dict[str, int] = {}

    for index, doc in enumerate(docs):
        validate_source_storage_path(doc.storage_path)
        if not doc.id or not doc.filename:
            raise ValueError("Each selected source document requires id and filename")
        file_bytes = supabase.storage.from_('uploads').download(doc.storage_path)
        filename = doc.filename
        fname_lower = filename.lower()
        doc_label = f"{filename} [retained_source:{doc.id}]"

        def add_text(file_class: str, text: str, label: str) -> None:
            nonlocal combined_text
            if not text:
                skipped.append(label)
                return
            max_chars, _max_count = REUNDERWRITE_TEXT_BUDGET.get(file_class, (10000, 3))
            source_files.append(label)
            file_classes[label] = file_class
            combined_text += (
                f"\n\n=== {label} ({file_class}) ===\n"
                "DOCUMENT_ROLE: Retained original source document\n"
                f"SOURCE_DOCUMENT_ID: {doc.id}\n"
                f"STORAGE_PATH: {doc.storage_path}\n"
                f"ORIGINAL_RUN_ID: {doc.run_id or 'none'}\n"
                f"SOURCE_KIND: {doc.source_kind or 'unknown'}\n"
                f"CONTENT_TYPE: {doc.content_type or 'unknown'}\n\n"
                f"{text[:max_chars]}"
            )

        if fname_lower.endswith(".zip"):
            zip_path = os.path.join(work_dir, f"{index}_source.zip")
            with open(zip_path, "wb") as f:
                f.write(file_bytes)
            with zipfile.ZipFile(zip_path, "r") as zf:
                for entry_name in zf.namelist():
                    base_name = Path(entry_name).name
                    if not base_name or base_name.startswith("."):
                        continue
                    file_class = classify_file(base_name)
                    if file_class == "unknown":
                        skipped.append(base_name)
                        continue
                    _max_chars, max_count = REUNDERWRITE_TEXT_BUDGET.get(file_class, (10000, 3))
                    if class_counts.get(file_class, 0) >= max_count:
                        skipped.append(base_name)
                        continue
                    entry_path = os.path.join(work_dir, f"{index}_source_{_safe_local_name(base_name)}")
                    with open(entry_path, "wb") as f:
                        f.write(zf.read(entry_name))
                    text = _extract_file_text(entry_path, base_name)
                    if base_name.lower().endswith(".pdf") and is_pdf_scanned(text):
                        text = await extract_scanned_pdf_text(entry_path, file_class.replace("_", " "))
                    class_counts[file_class] = class_counts.get(file_class, 0) + 1
                    add_text(file_class, text, f"{base_name} [retained_source:{doc.id}]")
            continue

        file_class = classify_file(filename)
        if file_class == "unknown":
            skipped.append(filename)
            continue
        _max_chars, max_count = REUNDERWRITE_TEXT_BUDGET.get(file_class, (10000, 3))
        if class_counts.get(file_class, 0) >= max_count:
            skipped.append(filename)
            continue
        file_path = os.path.join(work_dir, f"{index}_source_{_safe_local_name(filename)}")
        with open(file_path, "wb") as f:
            f.write(file_bytes)
        text = _extract_file_text(file_path, filename)
        if fname_lower.endswith(".pdf") and is_pdf_scanned(text):
            text = await extract_scanned_pdf_text(file_path, file_class.replace("_", " "))
        class_counts[file_class] = class_counts.get(file_class, 0) + 1
        add_text(file_class, text, doc_label)

    return combined_text, source_files, file_classes, skipped

def _merge_present_values(
    base: Any,
    updates: Any,
    *,
    path: str = "",
    changed_paths: set[str],
    conflicts: list[dict[str, Any]],
) -> Any:
    if isinstance(base, dict) and isinstance(updates, dict):
        merged = copy.deepcopy(base)
        for key, value in updates.items():
            child_path = f"{path}.{key}" if path else str(key)
            if key.startswith("_") and not _is_present(value):
                continue
            merged[key] = _merge_present_values(
                merged.get(key),
                value,
                path=child_path,
                changed_paths=changed_paths,
                conflicts=conflicts,
            )
        return merged
    if not _is_present(updates):
        return copy.deepcopy(base)
    if base != updates:
        changed_paths.add(path)
        if _is_present(base):
            conflicts.append({"path": path, "old_value": base, "new_value": updates})
    return copy.deepcopy(updates)

def _changed_top_level_fields(changed_paths: set[str]) -> set[str]:
    mapping = {
        "financials.fy25.revenue": "revenue",
        "financials.fy25.ebitda": "ebitda",
        "financials.fy25.total_labour_cost": "payroll_labour_cost",
        "financials.asking_price": "asking_price",
        "centre.licensed_places": "licensed_places",
        "occupancy.current_month_pct": "current_occupancy_pct",
        "occupancy.latest_week_pct": "latest_week_occupancy_pct",
        "occupancy.avg_4wk_pct": "avg_4wk_occupancy_pct",
        "occupancy.avg_13wk_pct": "avg_13wk_occupancy_pct",
        "occupancy.avg_52wk_pct": "avg_52wk_occupancy_pct",
        "financials.fy25.rent_pa": "rent_pa",
    }
    changed: set[str] = set()
    for path, field in mapping.items():
        if any(candidate == path or candidate.startswith(f"{path}.") for candidate in changed_paths):
            changed.add(field)
    return changed

def annotate_base_snapshot_facts(workflow: dict[str, Any], changed_fields: set[str], run_id: str) -> None:
    base_label = "Base run snapshot (not re-verified by selected diligence documents)"
    for key in ("facts", "extracted_facts"):
        facts = workflow.get(key)
        if not isinstance(facts, list):
            continue
        for fact in facts:
            if not isinstance(fact, dict):
                continue
            field = str(fact.get("field") or "")
            source = fact.get("source") if isinstance(fact.get("source"), dict) else {}
            if field not in changed_fields and not source.get("excerpt"):
                source["label"] = base_label
                source["file"] = base_label
                source["run_id"] = run_id
                fact["source"] = source
                fact["source_label"] = base_label
                fact["status"] = fact.get("status") or "needs_review"
    evidence = workflow.get("evidence")
    if isinstance(evidence, list):
        fact_fields = {
            str(f.get("evidence_id")): str(f.get("field") or "")
            for f in workflow.get("facts", []) if isinstance(f, dict)
        }
        for ev in evidence:
            if not isinstance(ev, dict):
                continue
            field = str(ev.get("field") or fact_fields.get(str(ev.get("id"))) or "")
            source = ev.get("source") if isinstance(ev.get("source"), dict) else {}
            if field not in changed_fields and not source.get("excerpt"):
                source["label"] = base_label
                source["file"] = base_label
                source["run_id"] = run_id
                ev["source"] = source
                ev["source_label"] = base_label
            else:
                source["run_id"] = run_id
                ev["source"] = source

def annotate_run_evidence_metadata(workflow: dict[str, Any], run_id: str) -> None:
    def scoped(local_id: Any) -> str | None:
        if not local_id:
            return None
        local = str(local_id)
        return local if ":" in local else f"{run_id}:{local}"

    for ev in workflow.get("evidence", []) or []:
        if not isinstance(ev, dict):
            continue
        local_id = str(ev.get("local_evidence_id") or ev.get("id") or "")
        if not local_id:
            continue
        ev["local_evidence_id"] = local_id
        ev["run_id"] = run_id
        ev["run_evidence_id"] = scoped(local_id)
        source = ev.get("source") if isinstance(ev.get("source"), dict) else {}
        source["local_evidence_id"] = local_id
        source["run_id"] = run_id
        source["run_evidence_id"] = scoped(local_id)
        ev["source"] = source

    for key in ("facts", "extracted_facts"):
        for fact in workflow.get(key, []) or []:
            if not isinstance(fact, dict):
                continue
            local_id = str(fact.get("local_evidence_id") or fact.get("evidence_id") or "")
            if not local_id:
                continue
            fact["local_evidence_id"] = local_id
            fact["run_id"] = run_id
            fact["run_evidence_id"] = scoped(local_id)
            source = fact.get("source") if isinstance(fact.get("source"), dict) else {}
            source["local_evidence_id"] = local_id
            source["run_id"] = run_id
            source["run_evidence_id"] = scoped(local_id)
            fact["source"] = source

def recalculate_total_score(scored: dict[str, Any]) -> None:
    dims = scored.get("dimensions", {})
    weighted_sum = 0.0
    weight_used = 0.0
    for dim_id, weight in DIMENSION_WEIGHTS.items():
        dim = dims.get(dim_id, {}) if isinstance(dims, dict) else {}
        raw = dim.get("score") if isinstance(dim, dict) else None
        if isinstance(raw, (int, float)) and 0 <= raw <= 10:
            weighted_sum += raw * weight
            weight_used += weight
    if weight_used > 0:
        scored["total_score"] = round((weighted_sum / weight_used) * 10, 1)

def reunderwrite_error(category: str, detail: str, status_code: int = 500) -> JSONResponse:
    return JSONResponse({"detail": detail, "error_category": category}, status_code=status_code)

def declared_total_bytes(req: ReunderwriteRequest) -> int:
    if isinstance(req.input_total_bytes, int) and req.input_total_bytes > 0:
        return req.input_total_bytes
    total = 0
    for doc in (req.selected_diligence_documents or []):
        if isinstance(doc.file_size, int) and doc.file_size > 0:
            total += doc.file_size
    for doc in (req.selected_source_documents or []):
        if isinstance(doc.file_size, int) and doc.file_size > 0:
            total += doc.file_size
    return total

@app.post("/pipeline/reunderwrite")
async def pipeline_reunderwrite(req: ReunderwriteRequest):
    work_dir = tempfile.mkdtemp(prefix="acquira-reunderwrite-")
    try:
        logger.info("reunderwrite.start run_id=%s deal_id=%s", req.run_id, req.deal_id)
        if req.mode != "reunderwrite":
            return reunderwrite_error("invalid_input", "mode must be reunderwrite", 400)
        if not req.deal_id or not req.run_id or not req.base_run_id:
            return reunderwrite_error("invalid_input", "deal_id, run_id, and base_run_id are required", 400)
        selected_diligence_documents = req.selected_diligence_documents or []
        selected_source_documents = req.selected_source_documents or []
        if not selected_diligence_documents and not selected_source_documents:
            return reunderwrite_error("invalid_input", "At least one selected diligence or source document is required", 400)
        total_docs = len(selected_diligence_documents) + len(selected_source_documents)
        total_bytes = declared_total_bytes(req)
        logger.info(
            "reunderwrite.inputs run_id=%s deal_id=%s source_docs=%s diligence_docs=%s declared_bytes=%s",
            req.run_id,
            req.deal_id,
            len(selected_source_documents),
            len(selected_diligence_documents),
            total_bytes,
        )
        if total_docs > MAX_REUNDERWRITE_DOCUMENTS:
            return reunderwrite_error("invalid_input", f"Select {MAX_REUNDERWRITE_DOCUMENTS} or fewer documents for one re-underwrite run.", 400)
        if total_bytes > MAX_REUNDERWRITE_DECLARED_BYTES:
            return reunderwrite_error("invalid_input", "Selected documents are too large for one synchronous re-underwrite run.", 400)

        base_extracted = copy.deepcopy(req.base.extracted)
        base_scored = copy.deepcopy(req.base.scored)
        base_workflow = copy.deepcopy(req.base.workflow)
        if not isinstance(base_extracted, dict) or not isinstance(base_scored, dict) or not isinstance(base_workflow, dict):
            return reunderwrite_error("invalid_input", "base.extracted, base.scored, and base.workflow must be objects", 400)

        try:
            source_text, source_source_files, source_file_classes, source_skipped = await parse_source_documents(
                selected_source_documents,
                work_dir=work_dir,
            ) if selected_source_documents else ("", [], {}, [])
            diligence_text, diligence_source_files, diligence_file_classes, diligence_skipped = await parse_diligence_documents(
                selected_diligence_documents,
                deal_id=req.deal_id,
                work_dir=work_dir,
            ) if selected_diligence_documents else ("", [], {}, [])
        except ValueError as e:
            logger.warning("reunderwrite.parse_invalid run_id=%s deal_id=%s error=%s", req.run_id, req.deal_id, e)
            return reunderwrite_error("invalid_input", str(e), 400)
        except Exception as e:
            logger.exception("reunderwrite.parse_failed run_id=%s deal_id=%s", req.run_id, req.deal_id)
            names = [doc.filename for doc in [*selected_source_documents, *selected_diligence_documents] if doc.filename]
            name_hint = ", ".join(names[:5])
            suffix = f" while processing {name_hint}" if name_hint else ""
            return reunderwrite_error("parse_failed", f"Selected document parse failed{suffix}: {e}", 422)

        combined_text = source_text + diligence_text
        source_files = [*source_source_files, *diligence_source_files]
        file_classes = {**source_file_classes, **diligence_file_classes}
        skipped = [*source_skipped, *diligence_skipped]
        if not combined_text.strip():
            return reunderwrite_error("parse_failed", "Could not extract text from selected documents.", 422)
        logger.info(
            "reunderwrite.parse_complete run_id=%s deal_id=%s parsed_files=%s skipped=%s",
            req.run_id,
            req.deal_id,
            len(source_files),
            len(skipped),
        )

        try:
            extraction_response = client.messages.create(
                model=MODEL,
                max_tokens=MAX_TOKENS,
                temperature=0,
                system=EXTRACTION_SYSTEM_PROMPT,
                messages=[{
                    "role": "user",
                    "content": (
                        "You are re-underwriting an existing childcare acquisition from selected retained source and diligence evidence.\n"
                        "Extract ONLY facts directly supported by the selected documents below. "
                        "Use the base snapshot only as context for entity matching and field names. "
                        "If a value is not present in the selected documents, return null for that field; "
                        "do not copy values from the base snapshot.\n\n"
                        f"BASE SNAPSHOT CONTEXT (not new evidence):\n{json.dumps(base_extracted, indent=2)[:40000]}\n\n"
                        f"SELECTED EVIDENCE:\n{combined_text[:CLAUDE_CHAR_LIMIT]}"
                    )
                }]
            )
            extracted_text = clean_json(extraction_response.content[0].text)
            diligence_extracted = json.loads(extracted_text)
        except Exception as e:
            logger.exception("reunderwrite.extraction_failed run_id=%s deal_id=%s", req.run_id, req.deal_id)
            return reunderwrite_error("extraction_failed", f"Extraction failed during re-underwrite: {e}", 502)
        logger.info("reunderwrite.extraction_complete run_id=%s deal_id=%s", req.run_id, req.deal_id)

        changed_paths: set[str] = set()
        merge_conflicts: list[dict[str, Any]] = []
        merged_extracted = _merge_present_values(
            base_extracted,
            diligence_extracted,
            changed_paths=changed_paths,
            conflicts=merge_conflicts,
        )

        merged_meta = merged_extracted.setdefault("meta", {})
        if isinstance(merged_meta, dict):
            merged_meta["reunderwrite_run_id"] = req.run_id
            merged_meta["base_run_id"] = req.base_run_id
            merged_meta["reunderwrite_source_type"] = "selected_source_and_diligence_documents"
            merged_meta["selected_source_document_ids"] = [doc.id for doc in selected_source_documents]
            merged_meta["selected_diligence_document_ids"] = [doc.id for doc in selected_diligence_documents]
            merged_meta["selected_source_files"] = source_files

        if req.pipeline_projects is not None:
            pipeline_supply = build_pipeline_supply(req.pipeline_projects, None)
            merged_extracted["_pipeline_projects"] = pipeline_supply["pipeline_projects"]
            merged_extracted["_pipeline_audit"] = pipeline_supply["pipeline_audit"]
        else:
            merged_extracted["_pipeline_projects"] = base_workflow.get("pipeline_projects") or base_scored.get("pipeline_projects") or []
            merged_extracted["_pipeline_audit"] = base_workflow.get("pipeline_audit") or base_scored.get("pipeline_audit")

        # Existing deals do not retain original source documents, so retain base
        # market context unless a later phase adds source-retention/recomputation.
        for key in ("_demand_context", "_market_context", "_market_audit"):
            if not merged_extracted.get(key) and base_extracted.get(key):
                merged_extracted[key] = copy.deepcopy(base_extracted[key])

        centre_name = merged_extracted.get("centre", {}).get("name") or base_scored.get("centre_name") or "centre"
        demand_context = merged_extracted.get("_demand_context")
        demand_block = "  No ABS demand data available.\n"
        if isinstance(demand_context, dict):
            demand_block = (
                f"  EDR (adj kids/place): {demand_context.get('adj_kids_per_place', {}).get('mid')}\n"
                f"  Zone: {demand_context.get('zone')}\n"
                f"  Confidence: {demand_context.get('confidence')}\n"
                f"  ABS hit: {demand_context.get('abs_hit')}\n"
                f"  Demand trend: {demand_context.get('demand_trend')}\n"
            )

        pipeline_intel_context = ""
        if merged_extracted.get("_pipeline_projects"):
            pipeline_intel_context += "\n\nSTRUCTURED DA / PIPELINE PROJECTS (preserved or provided):\n"
            for project in (merged_extracted.get("_pipeline_projects") or [])[:10]:
                pipeline_intel_context += (
                    f"- {project.get('status', 'unknown')}: "
                    f"{project.get('name') or project.get('address') or 'Pipeline project'}"
                    f" · {project.get('proposed_places') or 'unknown'} places"
                    f" · confidence {project.get('confidence', 'low')}\n"
                )

        try:
            scoring_response = client.messages.create(
                model=MODEL,
                max_tokens=MAX_TOKENS,
                temperature=0,
                system=SCORING_SYSTEM_PROMPT,
                messages=[{
                    "role": "user",
                    "content": (
                        "Re-score this childcare centre acquisition using the merged re-underwriting data.\n\n"
                        "The merged data starts from the prior run and overlays only source-backed values "
                        "found in selected retained source or diligence documents. Be explicit where underwriting still depends "
                        "on unresolved evidence. Do not treat waived diligence items as evidence.\n\n"
                        "DEMAND DATA (use for occupancy_demand external score, ignore any document demand figures):\n"
                        + demand_block
                        + f"\nMERGED EXTRACTED DATA:\n{json.dumps(merged_extracted, indent=2)}"
                        + f"{pipeline_intel_context}"
                    )
                }]
            )
            scored_text = clean_json(scoring_response.content[0].text)
            new_scored = json.loads(scored_text)
        except Exception as e:
            logger.exception("reunderwrite.scoring_failed run_id=%s deal_id=%s", req.run_id, req.deal_id)
            return reunderwrite_error("scoring_failed", f"Scoring failed during re-underwrite: {e}", 502)
        logger.info("reunderwrite.scoring_complete run_id=%s deal_id=%s", req.run_id, req.deal_id)
        recalculate_total_score(new_scored)
        new_scored["scoring_version"] = "2.4-reunderwrite"
        new_scored["scoring_timestamp"] = datetime.now(timezone.utc).isoformat()
        new_scored["reunderwrite_run_id"] = req.run_id
        new_scored["base_run_id"] = req.base_run_id
        new_scored["pipeline_projects"] = merged_extracted.get("_pipeline_projects") or []
        new_scored["pipeline_audit"] = merged_extracted.get("_pipeline_audit")
        if merged_extracted.get("_demand_context"):
            new_scored["effective_demand_ratio"] = merged_extracted["_demand_context"].get("adj_kids_per_place", {}).get("mid")
            new_scored["demand_zone"] = merged_extracted["_demand_context"].get("zone")
            new_scored["demand_context"] = merged_extracted["_demand_context"]
            new_scored["market_context"] = merged_extracted.get("_market_context")
            new_scored["market_audit"] = merged_extracted.get("_market_audit")

        try:
            structured_intel = build_structured_deal_intelligence(
                extracted=merged_extracted,
                scored=new_scored,
                combined_text=combined_text,
                source_files=source_files,
                file_classes=file_classes,
            )
        except Exception as e:
            logger.exception("reunderwrite.workflow_failed run_id=%s deal_id=%s", req.run_id, req.deal_id)
            return reunderwrite_error("workflow_build_failed", f"Workflow build failed during re-underwrite: {e}", 500)
        structured_intel["run_id"] = req.run_id
        structured_intel["base_run_id"] = req.base_run_id
        structured_intel["reunderwrite_source"] = {
            "mode": "selected_source_and_diligence_documents",
            "source_document_ids": [doc.id for doc in selected_source_documents],
            "diligence_document_ids": [doc.id for doc in selected_diligence_documents],
            "skipped_files": skipped,
        }
        changed_fields = _changed_top_level_fields(changed_paths)
        annotate_base_snapshot_facts(structured_intel, changed_fields, req.run_id)
        annotate_run_evidence_metadata(structured_intel, req.run_id)

        warnings = []
        if skipped:
            warnings.append(f"Skipped {len(skipped)} selected file(s) or ZIP entries that could not be parsed.")
        if not changed_paths:
            warnings.append("Selected documents were parsed but did not update any extracted fields.")
        if selected_source_documents:
            warnings.append("Re-underwrite used selected retained original source documents plus the base run snapshot.")
        else:
            warnings.append("Re-underwrite used base run snapshot plus selected diligence documents; no retained original source documents were selected.")

        try:
            diff = build_run_diff(
                base_extracted=base_extracted,
                base_scored=base_scored,
                base_workflow=base_workflow,
                new_extracted=merged_extracted,
                new_scored=new_scored,
                new_workflow=structured_intel,
                merge_conflicts=merge_conflicts,
                warnings=warnings,
            )
        except Exception as e:
            logger.exception("reunderwrite.diff_failed run_id=%s deal_id=%s", req.run_id, req.deal_id)
            return reunderwrite_error("diff_failed", f"Diff generation failed during re-underwrite: {e}", 500)
        logger.info("reunderwrite.diff_complete run_id=%s deal_id=%s", req.run_id, req.deal_id)

        logger.info("reunderwrite.completed run_id=%s deal_id=%s", req.run_id, req.deal_id)
        return {
            "run_id": req.run_id,
            "base_run_id": req.base_run_id,
            "status": "completed",
            "extracted": merged_extracted,
            "scored": new_scored,
            "workflow": structured_intel,
            "diff": diff,
            "meta": {
                "centre_name": centre_name,
                "source_files": source_files,
                "skipped_files": skipped,
            },
        }
    except ValueError as e:
        logger.warning("reunderwrite.invalid_input run_id=%s deal_id=%s error=%s", req.run_id, req.deal_id, e)
        return reunderwrite_error("invalid_input", str(e), 400)
    except Exception as e:
        logger.exception("reunderwrite.unexpected_failed run_id=%s deal_id=%s", req.run_id, req.deal_id)
        return reunderwrite_error("unexpected_failed", str(e), 500)
    finally:
        shutil.rmtree(work_dir, ignore_errors=True)


# ── P4.1: ACECQA Nearby Centres ───────────────────────────────────────────────
# TODO: ACECQA does not have a public REST API as of 2026.
# When one becomes available, replace this stub with a real integration.
# For now: accepts optional manual entry via query params, returns structured mock.
@app.get("/acecqa/nearby")
async def acecqa_nearby(
    lat: float = None,
    lng: float = None,
    radius_km: float = 2.0,
    postcode: str = None,
):
    """
    Returns nearby ACECQA-registered childcare centres within radius_km.
    TODO: Automate via ACECQA public API when available.
    Current implementation returns a structured mock for UI development.
    """
    return {
        "source": "mock",
        "note": "ACECQA does not provide a public API as of 2026. Integrate with NQS IT system when available.",
        "query": {"lat": lat, "lng": lng, "radius_km": radius_km, "postcode": postcode},
        "centres": [
            {
                "name": "Sunshine Early Learning Centre",
                "address": "45 Park St",
                "suburb": postcode or "Unknown",
                "nqs_rating": "Meeting NQS",
                "licensed_places": 60,
                "distance_km": 0.8,
                "lat": lat,
                "lng": lng,
                "provider": "Private",
            },
            {
                "name": "Little Stars Childcare",
                "address": "12 Main Rd",
                "suburb": postcode or "Unknown",
                "nqs_rating": "Exceeding NQS",
                "licensed_places": 80,
                "distance_km": 1.4,
                "lat": lat,
                "lng": lng,
                "provider": "Community",
            },
            {
                "name": "Rainbow Kids Centre",
                "address": "7 Station Ave",
                "suburb": postcode or "Unknown",
                "nqs_rating": "Working Towards NQS",
                "licensed_places": 45,
                "distance_km": 1.9,
                "lat": lat,
                "lng": lng,
                "provider": "Private",
            },
        ],
    }


# ── P4.2: ABS Demographic Lookup ─────────────────────────────────────────────
# TODO: ABS Data API (https://api.data.abs.gov.au) is available but requires
# knowing the correct dataset IDs for small-area population data.
# The correct dataset for 0-4 population by SA2 is ABS_CENSUS2021_B04.
# When properly integrated, replace mock below with real ABS API call.
import urllib.request
import urllib.error

@app.get("/demographics/{postcode}")
async def demographics(postcode: str):
    """
    Returns catchment demographic data for a postcode.
    Attempts ABS Data API; falls back to structured mock on failure.
    TODO: Implement full SA2-level 0-4 population trend from ABS Census data.
    """
    # Attempt ABS API
    try:
        abs_url = f"https://api.data.abs.gov.au/data/ABS_CENSUS2021_B04/1+2+3..?startPeriod=2021&dimensionAtObservation=AllDimensions&format=jsondata"
        req = urllib.request.Request(abs_url, headers={"Accept": "application/json"}, method="GET")
        # Short timeout — fall back to mock if ABS is slow
        with urllib.request.urlopen(req, timeout=3) as resp:
            # TODO: parse the correct 0-4 age group for the postcode's SA2
            pass  # Fall through to mock for now
    except Exception:
        pass  # Expected — fall back to mock

    # Structured mock with realistic shape
    return {
        "source": "mock",
        "note": "TODO: Integrate ABS Census 2021 SA2 population data for postcode-to-SA2 mapping. Dataset: ABS_CENSUS2021_B04.",
        "postcode": postcode,
        "population_0_4": {
            "2016": 412,
            "2021": 385,
            "trend": "declining",
            "pct_change_5yr": -6.6,
            "risk_flag": True,
            "risk_note": "0-4 population declined 6.6% over 5 years — potential demand headwind.",
        },
        "population_total": {
            "2016": 8200,
            "2021": 8650,
            "trend": "stable",
        },
        "median_household_income": 104000,
        "dual_income_pct_estimated": 62,
        "ccs_eligibility_estimated": "high",
    }

# ── P6: DA Pipeline helpers ───────────────────────────────────────────────────

CHILDCARE_KEYWORDS = [
    "child care", "childcare", "early learning", "early childhood",
    "education centre", "kindergarten", "kinder"
]

def extract_places_from_description(description: str) -> Optional[int]:
    """Extract a number of licensed places from a DA description."""
    patterns = [
        r'(\d+)\s*(?:child\s*)?places',
        r'(\d+)\s*children',
        r'(\d+)\s*child\s*places',
        r'(\d+)[- ]place',
    ]
    for pat in patterns:
        m = re.search(pat, description, re.I)
        if m:
            return int(m.group(1))
    return None

def classify_da_status(description: str) -> str:
    """Classify a DA status from its description."""
    desc_lower = description.lower()
    if any(w in desc_lower for w in ["approved", "granted", "permit issued", "development approval"]):
        return "approved"
    if any(w in desc_lower for w in ["refused", "rejected", "not approved"]):
        return "refused"
    if any(w in desc_lower for w in ["lodged", "submitted", "under assessment", "application received"]):
        return "lodged"
    return "unknown"

def assess_pipeline_risk(applications: list, existing_licensed_places: int) -> dict:
    """Assess supply risk from DA pipeline."""
    approved_places = sum(
        a.get("places") or 0
        for a in applications
        if a.get("status") == "approved" and a.get("places")
    )
    if existing_licensed_places > 0:
        ratio = approved_places / existing_licensed_places
    else:
        ratio = 0

    if ratio > 0.5:
        risk_level = "HIGH"
    elif ratio > 0.25:
        risk_level = "MEDIUM"
    else:
        risk_level = "LOW"

    return {
        "risk_level": risk_level,
        "approved_pipeline_places": approved_places,
        "ratio_to_licensed": round(ratio, 2),
    }


@app.get("/planning/nearby")
async def planning_nearby(
    postcode: str = Query(..., description="Postcode to search"),
    suburb: Optional[str] = Query(None),
    state: str = Query("VIC"),
    radius_km: float = Query(2.0),
):
    """
    Returns nearby DA applications for childcare centres.
    Uses PlanningAlerts API if PLANNING_ALERTS_API_KEY is set, otherwise mock data.
    """
    api_key = os.environ.get("PLANNING_ALERTS_API_KEY")

    if api_key and suburb:
        # Live data from PlanningAlerts
        try:
            import urllib.request as ur
            url = (
                f"https://api.planningalerts.org.au/applications.js"
                f"?key={api_key}&suburb={suburb}&state={state}&count=200"
            )
            req = ur.Request(url, headers={"Accept": "application/json"}, method="GET")
            with ur.urlopen(req, timeout=10) as resp:
                raw = json.loads(resp.read().decode())

            applications = []
            for item in (raw if isinstance(raw, list) else raw.get("applications", [])):
                desc = item.get("description") or ""
                # Filter for childcare keywords
                if not any(kw in desc.lower() for kw in CHILDCARE_KEYWORDS):
                    continue
                places = extract_places_from_description(desc)
                status = classify_da_status(desc)
                applications.append({
                    "address":       item.get("address"),
                    "description":   desc,
                    "status":        status,
                    "date":          item.get("date_received") or item.get("date_scraped"),
                    "places":        places,
                    "distance_km":   None,
                    "info_url":      item.get("info_url"),
                    "on_notice_from": item.get("on_notice_from"),
                    "on_notice_to":   item.get("on_notice_to"),
                })

            approved = [a for a in applications if a["status"] == "approved"]
            lodged   = [a for a in applications if a["status"] == "lodged"]
            refused  = [a for a in applications if a["status"] == "refused"]
            total_approved_places  = sum(a["places"] or 0 for a in approved)
            total_pipeline_places  = sum(a["places"] or 0 for a in lodged)
            risk_flag = total_approved_places > 0
            risk_note = (
                f"{len(approved)} approved DA{'s' if len(approved) != 1 else ''} add "
                f"{total_approved_places} pipeline places within search area"
                if risk_flag else "No approved childcare DAs found in search area"
            )

            return {
                "source": "live",
                "postcode": postcode,
                "suburb": suburb,
                "state": state,
                "applications": applications,
                "summary": {
                    "total": len(applications),
                    "approved": len(approved),
                    "lodged": len(lodged),
                    "refused": len(refused),
                    "total_approved_places": total_approved_places,
                    "total_pipeline_places": total_pipeline_places,
                    "risk_flag": risk_flag,
                    "risk_note": risk_note,
                }
            }
        except Exception as e:
            print(f"[planning/nearby] Live fetch failed: {e}. Falling back to mock.")

    # Mock data
    suburb_display = suburb or "Forest Hill"
    return {
        "source": "mock",
        "note": "Live DA data requires PlanningAlerts API key. Showing illustrative examples only.",
        "postcode": postcode,
        "suburb": suburb_display,
        "state": state,
        "applications": [
            {
                "address": f"45 Example St, {suburb_display} {state} {postcode}",
                "description": "Construction of a child care centre (90 places)",
                "status": "approved",
                "date": "2025-08-14",
                "places": 90,
                "distance_km": 1.2,
                "info_url": None,
            },
            {
                "address": f"12 Sample Ave, Nearby Suburb {state} {postcode}",
                "description": "Early learning centre - 75 places",
                "status": "lodged",
                "date": "2025-11-22",
                "places": 75,
                "distance_km": 1.8,
                "info_url": None,
            },
            {
                "address": f"88 Demo Rd, Adjacent Suburb {state} {postcode}",
                "description": "Child care centre expansion - additional 30 places",
                "status": "approved",
                "date": "2025-06-03",
                "places": 30,
                "distance_km": 2.1,
                "info_url": None,
            },
        ],
        "summary": {
            "total": 3,
            "approved": 2,
            "lodged": 1,
            "refused": 0,
            "total_approved_places": 120,
            "total_pipeline_places": 75,
            "risk_flag": True,
            "risk_note": f"2 approved DAs add 120 pipeline places within {radius_km}km — significant supply risk",
        }
    }


@app.get("/planning/councils")
async def planning_councils():
    """
    Returns a static reference list of major Australian councils with their planning portal URLs.
    Useful for manual DA research.
    """
    return {
        "source": "static",
        "last_updated": "2026-03",
        "councils": {
            "VIC": [
                {"name": "Melbourne City Council",        "url": "https://development.melbourne.vic.gov.au/planning-register", "notes": "Search: child care, early learning"},
                {"name": "Boroondara City Council",       "url": "https://eservices.boroondara.vic.gov.au/datrack/",          "notes": "Search: childcare, early childhood"},
                {"name": "Monash City Council",           "url": "https://www.monash.vic.gov.au/Planning-Building/Planning/Planning-Applications", "notes": "Search: child care centre"},
                {"name": "Whitehorse City Council",       "url": "https://www.whitehorse.vic.gov.au/planning-applications",   "notes": "Search: early learning, childcare"},
                {"name": "Knox City Council",             "url": "https://www.knox.vic.gov.au/planning-permits",              "notes": "Search: child care"},
                {"name": "Manningham City Council",       "url": "https://www.manningham.vic.gov.au/building-planning/planning/planning-applications", "notes": "Search: early learning centre"},
                {"name": "Maroondah City Council",        "url": "https://www.maroondah.vic.gov.au/Planning-permits",         "notes": "Search: childcare, kinder"},
                {"name": "Yarra City Council",            "url": "https://www.yarracity.vic.gov.au/planning-and-building/planning-applications", "notes": "Search: child care, early childhood"},
                {"name": "Glen Eira City Council",        "url": "https://www.gleneira.vic.gov.au/planning-permits",          "notes": "Search: child care centre"},
                {"name": "Bayside City Council",          "url": "https://www.bayside.vic.gov.au/planning",                  "notes": "Search: early learning, childcare"},
                {"name": "Moonee Valley City Council",    "url": "https://www.mvcc.vic.gov.au/planning",                     "notes": "Search: child care"},
                {"name": "Darebin City Council",          "url": "https://www.darebin.vic.gov.au/planning-building-permits",  "notes": "Search: early childhood, childcare"},
            ],
            "NSW": [
                {"name": "City of Sydney Council",        "url": "https://da.cityofsydney.nsw.gov.au/",                      "notes": "Search: child care centre, early learning"},
                {"name": "Northern Beaches Council",      "url": "https://www.northernbeaches.nsw.gov.au/services/planning-and-building/development-applications", "notes": "Search: child care"},
                {"name": "Ku-ring-gai Council",           "url": "https://www.kmc.nsw.gov.au/planning_and_development/development_applications", "notes": "Search: early learning, childcare"},
                {"name": "Ryde City Council",             "url": "https://www.ryde.nsw.gov.au/Planning/Development-Applications", "notes": "Search: child care centre"},
                {"name": "Parramatta City Council",       "url": "https://eservices.parracity.nsw.gov.au/datracking/",        "notes": "Search: early childhood, childcare"},
                {"name": "Blacktown City Council",        "url": "https://www.blacktown.nsw.gov.au/Planning-and-Building/Development-applications", "notes": "Search: child care"},
                {"name": "Lane Cove Council",             "url": "https://www.lanecove.nsw.gov.au/planning/development-applications/", "notes": "Search: early learning"},
                {"name": "Willoughby City Council",       "url": "https://www.willoughby.nsw.gov.au/Planning-Building/Development-Applications", "notes": "Search: child care, kinder"},
                {"name": "Mosman Council",                "url": "https://www.mosman.nsw.gov.au/council/services/planning/development-applications", "notes": "Search: early childhood"},
                {"name": "Strathfield Council",           "url": "https://www.strathfield.nsw.gov.au/building-planning/development-applications", "notes": "Search: childcare, early learning"},
                {"name": "Canterbury-Bankstown Council",  "url": "https://www.cbcity.nsw.gov.au/planning-and-building/development-applications", "notes": "Search: child care centre"},
            ],
            "QLD": [
                {"name": "Brisbane City Council",         "url": "https://developmenti.brisbane.qld.gov.au/",                "notes": "Search: child care, early learning"},
                {"name": "Gold Coast City Council",       "url": "https://eplanning.goldcoast.qld.gov.au/",                  "notes": "Search: childcare, early childhood"},
                {"name": "Sunshine Coast Council",        "url": "https://eplanning.sunshinecoast.qld.gov.au/",              "notes": "Search: child care centre"},
                {"name": "Moreton Bay Regional Council",  "url": "https://eplanning.moretonbay.qld.gov.au/",                 "notes": "Search: early learning, childcare"},
                {"name": "Logan City Council",            "url": "https://eplanning.logan.qld.gov.au/",                      "notes": "Search: child care"},
                {"name": "Ipswich City Council",          "url": "https://eplanning.ipswich.qld.gov.au/",                    "notes": "Search: early childhood centre"},
                {"name": "Townsville City Council",       "url": "https://eplanning.townsville.qld.gov.au/",                 "notes": "Search: childcare, child care"},
                {"name": "Cairns Regional Council",       "url": "https://eplanning.cairns.qld.gov.au/",                     "notes": "Search: early learning"},
                {"name": "Redland City Council",          "url": "https://eplanning.redland.qld.gov.au/",                    "notes": "Search: child care centre"},
                {"name": "Toowoomba Regional Council",    "url": "https://eplanning.toowoomba.qld.gov.au/",                  "notes": "Search: childcare, kinder"},
                {"name": "Rockhampton Regional Council",  "url": "https://eplanning.rockhamptonregion.qld.gov.au/",          "notes": "Search: child care"},
            ],
            "WA": [
                {"name": "City of Perth",                 "url": "https://www.perth.wa.gov.au/planning-development/development-applications", "notes": "Search: child care, early learning"},
                {"name": "City of Stirling",              "url": "https://www.stirling.wa.gov.au/planning",                  "notes": "Search: childcare centre"},
                {"name": "City of Joondalup",             "url": "https://www.joondalup.wa.gov.au/planning",                 "notes": "Search: early learning, child care"},
                {"name": "City of Swan",                  "url": "https://www.swan.wa.gov.au/planning",                     "notes": "Search: childcare"},
                {"name": "City of Melville",              "url": "https://www.melvillecity.com.au/planning",                 "notes": "Search: child care centre"},
                {"name": "City of Canning",               "url": "https://www.canning.wa.gov.au/planning",                  "notes": "Search: early childhood"},
                {"name": "City of Gosnells",              "url": "https://www.gosnells.wa.gov.au/planning",                  "notes": "Search: childcare, child care"},
                {"name": "City of Wanneroo",              "url": "https://www.wanneroo.wa.gov.au/planning",                  "notes": "Search: early learning"},
            ],
            "SA": [
                {"name": "City of Adelaide",              "url": "https://www.cityofadelaide.com.au/planning",               "notes": "Search: child care, early learning"},
                {"name": "City of Charles Sturt",         "url": "https://www.charlessturt.sa.gov.au/planning",              "notes": "Search: childcare centre"},
                {"name": "City of Onkaparinga",           "url": "https://www.onkaparinga.sa.gov.au/planning",               "notes": "Search: early childhood, child care"},
                {"name": "City of Marion",                "url": "https://www.marion.sa.gov.au/planning",                    "notes": "Search: childcare"},
                {"name": "City of Tea Tree Gully",        "url": "https://www.teatreegully.sa.gov.au/planning",              "notes": "Search: early learning"},
                {"name": "City of Salisbury",             "url": "https://www.salisbury.sa.gov.au/planning",                 "notes": "Search: child care centre"},
                {"name": "City of Port Adelaide Enfield", "url": "https://www.portenf.sa.gov.au/planning",                  "notes": "Search: childcare, kinder"},
                {"name": "City of Prospect",              "url": "https://www.prospect.sa.gov.au/planning",                  "notes": "Search: early childhood"},
            ],
        }
    }


@app.post("/pipeline")
async def pipeline(req: PipelineRequest):
    """
    Streaming pipeline endpoint using Server-Sent Events.

    Accepts single file (legacy) or multiple files/ZIP (new).
    Event types:
      progress  — { step, label, detail? }
      error     — { message }
      complete  — { extracted, scored, meta }
    """

    async def generate():
        work_dir      = tempfile.mkdtemp(prefix='acquira-')
        storage_paths = req.resolved_paths()
        all_filenames = req.resolved_filenames()
        pipeline_request_id = uuid.uuid4().hex
        retained_source_files: list[dict[str, Any]] = []
        print("[pipeline] release:", API_RELEASE)

        try:
            # ── Step 1: Download & parse all files ───────────────────────
            yield sse_event("progress", {
                "step": 1, "total": 5,
                "label": "Downloading files",
                "detail": f"{len(storage_paths)} file{'s' if len(storage_paths) != 1 else ''}"
            })

            # Text budget per file class: (max_chars_per_file, max_file_count)
            # Keeps combined_text under ~120k chars even with 30 files.
            TEXT_BUDGET: dict[str, tuple[int, int]] = {
                'im_pdf':               (50000, 3),
                'im_docx':              (50000, 3),
                'pl_excel':             (10000, 10),
                'occupancy_excel':      (10000, 10),
                'transaction_excel':    (8000,  5),
                'payroll_excel':        (8000,  5),
                'lease_pdf':            (8000,  5),
                'lease_docx':           (8000,  5),
                'service_approval_pdf': (4000,  3),
                'nqs_pdf':              (4000,  3),
            }
            CLAUDE_CHAR_LIMIT = 120_000

            class_counts: dict[str, int] = {}
            combined_text = ''
            source_files:  list[str] = []
            file_classes:  dict[str, str] = {}
            skipped:       list[str] = []
            retention_warnings: list[str] = []

            for source_index, (storage_path, filename) in enumerate(zip(storage_paths, all_filenames)):
                fname_lower = filename.lower()
                try:
                    validate_pipeline_temp_path(storage_path)
                except Exception as e:
                    print("[pipeline] invalid Supabase storage path")
                    print("  bucket: uploads")
                    print("  storage_path:", storage_path)
                    print("  filename:", filename)
                    print("  error:", repr(e))

                    yield sse_event("error", {
                        "message": (
                            "Uploaded file path is invalid. "
                            f"bucket=uploads path={storage_path}. "
                            "Please re-upload and retry."
                        )
                    })
                    return

                try:
                    print("[pipeline] downloading file")
                    print("  bucket: uploads")
                    print("  storage_path:", storage_path)
                    print("  filename:", filename)
                    file_bytes = supabase.storage.from_('uploads').download(storage_path)
                except Exception as e:
                    print("[pipeline] Supabase download failed")
                    print("  bucket: uploads")
                    print("  storage_path:", storage_path)
                    print("  filename:", filename)
                    print("  error:", repr(e))

                    yield sse_event("error", {
                        "message": (
                            "Uploaded file could not be read from Supabase Storage. "
                            f"bucket=uploads path={storage_path}. "
                            "This usually means the frontend sent an invalid path or the file was deleted by a previous pipeline run. "
                            "Please re-upload and retry."
                        )
                    })
                    return

                try:
                    retained_source_files.append(retain_pipeline_source_file(
                        pipeline_request_id=pipeline_request_id,
                        index=source_index,
                        original_storage_path=storage_path,
                        filename=filename,
                        file_bytes=file_bytes,
                    ))
                except Exception as e:
                    warning = f"Source retention failed for {filename}; continuing with uploaded file for this run."
                    retention_warnings.append(warning)
                    print("[pipeline] Supabase source retention failed (non-fatal)")
                    print("  bucket: uploads")
                    print("  storage_path:", storage_path)
                    print("  filename:", filename)
                    print("  error:", repr(e))

                # ── ZIP: extract contained files ──────────────────────────
                if fname_lower.endswith('.zip'):
                    yield sse_event("progress", {
                        "step": 1, "total": 5,
                        "label": "Unpacking ZIP",
                        "detail": filename
                    })
                    zip_path = os.path.join(work_dir, 'upload.zip')
                    with open(zip_path, 'wb') as f:
                        f.write(file_bytes)

                    with zipfile.ZipFile(zip_path, 'r') as zf:
                        for entry_name in zf.namelist():
                            base_name = Path(entry_name).name
                            if not base_name or base_name.startswith('.'): continue

                            file_class = classify_file(base_name)
                            if file_class == 'unknown': continue

                            max_chars, max_count = TEXT_BUDGET.get(file_class, (4000, 3))
                            if class_counts.get(file_class, 0) >= max_count:
                                skipped.append(base_name)
                                continue

                            entry_path = os.path.join(
                                work_dir, re.sub(r'[^a-zA-Z0-9._-]', '_', base_name)
                            )
                            with open(entry_path, 'wb') as f:
                                f.write(zf.read(entry_name))

                            text = _extract_file_text(entry_path, base_name)

                            # Vision fallback for scanned PDFs
                            if base_name.lower().endswith('.pdf') and is_pdf_scanned(text):
                                yield sse_event("progress", {
                                    "step": 1, "total": 5,
                                    "label": "Reading scanned PDF",
                                    "detail": base_name
                                })
                                text = await extract_scanned_pdf_text(
                                    entry_path, file_class.replace('_', ' ')
                                )

                            if text:
                                combined_text += f'\n\n=== {base_name} ({file_class}) ===\n{text[:max_chars]}'
                                source_files.append(base_name)
                                file_classes[base_name] = file_class
                                class_counts[file_class] = class_counts.get(file_class, 0) + 1

                # ── Single file ───────────────────────────────────────────
                else:
                    file_class = classify_file(filename)
                    if file_class == 'unknown':
                        skipped.append(filename)
                        continue

                    max_chars, max_count = TEXT_BUDGET.get(file_class, (30000, 1))
                    if class_counts.get(file_class, 0) >= max_count:
                        skipped.append(filename)
                        continue

                    file_path = os.path.join(
                        work_dir, re.sub(r'[^a-zA-Z0-9._-]', '_', filename)
                    )
                    with open(file_path, 'wb') as f:
                        f.write(file_bytes)

                    text = _extract_file_text(file_path, filename)

                    # Vision fallback for scanned PDFs
                    if fname_lower.endswith('.pdf') and is_pdf_scanned(text):
                        yield sse_event("progress", {
                            "step": 1, "total": 5,
                            "label": "Reading scanned PDF",
                            "detail": "Using vision extraction"
                        })
                        text = await extract_scanned_pdf_text(
                            file_path, file_class.replace('_', ' ')
                        )

                    if text:
                        combined_text += f'\n\n=== {filename} ({file_class}) ===\n{text[:max_chars]}'
                        source_files.append(filename)
                        file_classes[filename] = file_class
                        class_counts[file_class] = class_counts.get(file_class, 0) + 1
                    else:
                        skipped.append(filename)

            if skipped:
                print(f"[pipeline] skipped {len(skipped)} file(s): {skipped}")

            if not combined_text.strip():
                yield sse_event("error", {"message": "Could not extract text from any uploaded file."})
                return

            # ── Step 2: Extract ───────────────────────────────────────────
            yield sse_event("progress", {
                "step": 2, "total": 5,
                "label": "Extracting metrics",
                "detail": f"Reading {len(source_files)} file{'s' if len(source_files) != 1 else ''} · {len(combined_text):,} characters"
            })

            extraction_response = client.messages.create(
                model=MODEL,
                max_tokens=MAX_TOKENS,
                temperature=0,
                system=EXTRACTION_SYSTEM_PROMPT,
                messages=[{
                    "role": "user",
                    "content": (
                        f"Extract structured data from this childcare centre document.\n\n"
                        f"Source files: {', '.join(source_files)}\n\n"
                        f"CONTENT:\n{combined_text[:CLAUDE_CHAR_LIMIT]}"
                    )
                }]
            )
            extracted_text = clean_json(extraction_response.content[0].text)
            extracted      = json.loads(extracted_text)
            centre_name    = extracted.get('centre', {}).get('name') or 'centre'

            # ── Step 2.5: Compute deterministic demand & market context ───────
            # Python is the authoritative scorer for market_position.
            # We resolve postcode and licensed places from extracted data,
            # compute EDR + market score, then inject into extracted before
            # passing to the LLM. LLM explains; it doesn't generate.
            try:
                _postcode = (
                    extracted.get("centre", {}).get("postcode")
                    or extracted.get("key_ratios", {}).get("postcode")
                    or ""
                )
                _postcode = str(_postcode).strip().zfill(4) if _postcode else ""

                _licensed_places = (
                    extracted.get("centre", {}).get("licensed_places")
                    or extracted.get("key_ratios", {}).get("licensed_places")
                    or 0
                )

                # Pipeline intel: structured DA projects plus legacy count support.
                # Legacy approved_das stays score-stable: each approved DA still
                # contributes 90 approved places, matching prior behavior.
                _pi = req.pipelineIntel
                _pipeline_supply = build_pipeline_supply(
                    req.pipelineProjects,
                    _pi,
                    search_radius_km=None,
                )
                _pipeline_projects = _pipeline_supply["pipeline_projects"]
                _pipeline_audit = _pipeline_supply["pipeline_audit"]
                _approved_pipeline_places = _pipeline_audit["approved_places"]
                _pipeline_source = (
                    "Structured manual DA projects."
                    if _pipeline_audit["source_type"] == "manual_structured"
                    else "Legacy user-provided DA count; approved DAs assume 90 places each."
                    if _pipeline_audit["source_type"] == "manual_legacy_count"
                    else None
                )
                _pipeline_searched = bool(_pipeline_audit["searched"])

                # Count competitors from pipeline_mentions as proxy when no explicit intel
                _competitor_count = 0  # will be refined by LLM narrative
                _competitor_source = "Not verified"
                _included_centres = 0

                if _postcode:
                    _subject_places = _licensed_places or 60
                    _centre = extracted.get("centre", {}) or {}
                    _postcode_supply = {
                        "source": "unavailable",
                        "confidence": "low",
                        "total_licensed_places": _subject_places,
                        "competitor_count": 0,
                        "included_centres": 1,
                        "source_label": "Subject licensed places fallback / manual estimate",
                        "warnings": [],
                    }
                    try:
                        _acecqa_resp = supabase.from_("acecqa_centres") \
                            .select("licensed_places") \
                            .eq("postcode", _postcode) \
                            .execute()
                        _acecqa_rows = _acecqa_resp.data or []
                        _acecqa_total = sum((r.get("licensed_places") or 0) for r in _acecqa_rows)
                        if _acecqa_total > 0:
                            _postcode_supply = {
                                "source": "postcode_fallback",
                                "confidence": "medium",
                                "total_licensed_places": _acecqa_total + _subject_places,
                                "competitor_count": max(len(_acecqa_rows) - 1, 0),
                                "included_centres": len(_acecqa_rows) + 1,
                                "source_label": "ACECQA centres table filtered by postcode plus subject centre.",
                                "warnings": [],
                            }
                        else:
                            _is_reg = POSTCODE_AREA_KM2.get(_postcode, 50) > 200
                            _share = 0.25 if _is_reg else 0.12
                            _pipeline_mentions = extracted.get("pipeline_mentions", []) or []
                            _postcode_supply = {
                                "source": "market_share_heuristic",
                                "confidence": "low",
                                "total_licensed_places": round(_subject_places / _share),
                                "competitor_count": max(len(_pipeline_mentions), 0),
                                "included_centres": max(len(_pipeline_mentions), 0) + 1 if _pipeline_mentions else 1,
                                "source_label": "Fallback market-share heuristic; competitor count from document pipeline mentions if present.",
                                "warnings": ["Postcode ACECQA rows were unavailable; used market-share heuristic."],
                            }
                    except Exception as _sq_err:
                        print(f"[demand_service] Supabase ACECQA query failed (non-fatal): {_sq_err}")
                        _pipeline_mentions = extracted.get("pipeline_mentions", []) or []
                        _postcode_supply = {
                            "source": "postcode_query_failed",
                            "confidence": "low",
                            "total_licensed_places": _subject_places,
                            "competitor_count": max(len(_pipeline_mentions), 0),
                            "included_centres": max(len(_pipeline_mentions), 0) + 1 if _pipeline_mentions else 1,
                            "source_label": "Fallback after ACECQA query failure; competitor count from document pipeline mentions if present.",
                            "warnings": [f"Postcode ACECQA query failed: {_sq_err}"],
                        }

                    _radius_km = POSTCODE_AREA_KM2.get(_postcode)
                    _geo_supply = get_nearby_competitors(
                        supabase,
                        address=_centre.get("address"),
                        suburb=_centre.get("suburb"),
                        state=_centre.get("state"),
                        postcode=_postcode,
                        radius_km=(
                            2.0 if _radius_km and _radius_km < 20
                            else 3.0 if not _radius_km or _radius_km < 80
                            else 5.0
                        ),
                        subject_name=_centre.get("name"),
                        subject_address=_centre.get("address"),
                        service_approval_number=_centre.get("service_approval_number"),
                        subject_licensed_places=_subject_places,
                        lat=_centre.get("lat") or _centre.get("latitude"),
                        lng=_centre.get("lng") or _centre.get("longitude"),
                    )
                    _use_geo_for_scoring = (
                        _geo_supply.get("source") == "geospatial_supabase"
                        and _geo_supply.get("confidence") == "high"
                    )
                    _scoring_supply = _geo_supply if _use_geo_for_scoring else _postcode_supply
                    _total_catchment_places = _scoring_supply.get("total_licensed_places") or _subject_places
                    _competitor_count = _scoring_supply.get("competitor_count") or 0
                    _included_centres = (
                        _competitor_count + 1
                        if _scoring_supply.get("source") == "geospatial_supabase"
                        else _scoring_supply.get("included_centres") or (_competitor_count + 1 if _competitor_count else 1)
                    )
                    _competitor_source = (
                        "ACECQA geospatial radius via Supabase RPC get_nearby_centres."
                        if _scoring_supply.get("source") == "geospatial_supabase"
                        else _postcode_supply.get("source_label")
                    )

                    _demand_ctx = compute_demand(_postcode, _total_catchment_places)
                    _pipeline_audit["search_radius_km"] = _demand_ctx.get("radius_km")
                    _market_ctx = market_position_score(
                        demand_context=_demand_ctx,
                        competitor_count=_competitor_count,
                        approved_pipeline_places=_approved_pipeline_places,
                        subject_licensed_places=_subject_places,
                    )
                    _postcode_demand_ctx = compute_demand(_postcode, _postcode_supply.get("total_licensed_places") or _subject_places)
                    _geo_demand_ctx = (
                        compute_demand(_postcode, _geo_supply.get("total_licensed_places") or _subject_places)
                        if _geo_supply.get("source") == "geospatial_supabase"
                        else None
                    )
                    _material_difference = material_supply_difference(
                        geospatial=_geo_supply,
                        postcode_supply=_postcode_supply,
                        geospatial_edr=(_geo_demand_ctx or {}).get("adj_kids_per_place", {}).get("mid"),
                        postcode_edr=_postcode_demand_ctx.get("adj_kids_per_place", {}).get("mid"),
                        geospatial_zone=(_geo_demand_ctx or {}).get("zone"),
                        postcode_zone=_postcode_demand_ctx.get("zone"),
                    )
                    _competitor_supply = {
                        "source": _geo_supply.get("source"),
                        "confidence": _geo_supply.get("confidence"),
                        "radius_km": _geo_supply.get("radius_km"),
                        "competitor_count": _geo_supply.get("competitor_count"),
                        "total_licensed_places": _geo_supply.get("total_licensed_places"),
                        "target_geocode_method": _geo_supply.get("target_geocode_method"),
                        "exclusion_method": _geo_supply.get("exclusion_method"),
                        "scoring_source": _scoring_supply.get("source"),
                        "scoring_confidence": _scoring_supply.get("confidence"),
                        "compared_to_postcode": {
                            "competitor_count": _postcode_supply.get("competitor_count"),
                            "total_licensed_places": _postcode_supply.get("total_licensed_places"),
                            "edr": _postcode_demand_ctx.get("adj_kids_per_place", {}).get("mid"),
                        },
                        "material_difference": _material_difference,
                        "warnings": [
                            *(_geo_supply.get("warnings") or []),
                            *(_postcode_supply.get("warnings") or []),
                        ],
                    }
                    _vendor_kids = None
                    for _path in (
                        ("market", "kids_0_4"),
                        ("market", "children_0_4"),
                        ("demographics", "kids_0_4"),
                        ("demographics", "children_0_4"),
                        ("key_ratios", "kids_0_4"),
                    ):
                        _cur = extracted
                        for _key in _path:
                            _cur = _cur.get(_key, {}) if isinstance(_cur, dict) else {}
                        if isinstance(_cur, (int, float)):
                            _vendor_kids = int(_cur)
                            break
                    _market_audit = build_market_audit(
                        _demand_ctx,
                        _market_ctx,
                        subject_licensed_places=_subject_places,
                        competitor_source=_competitor_source,
                        included_centres=_included_centres,
                        pipeline_source=_pipeline_source,
                        pipeline_searched=_pipeline_searched,
                        pipeline_audit=_pipeline_audit,
                        vendor_kids_0_4=_vendor_kids,
                        competitor_supply=_competitor_supply,
                    )
                    extracted["_demand_context"] = _demand_ctx
                    extracted["_market_context"] = _market_ctx
                    extracted["_market_audit"] = _market_audit
                    extracted["_pipeline_projects"] = _pipeline_projects
                    extracted["_pipeline_audit"] = _pipeline_audit
                else:
                    extracted["_demand_context"] = None
                    extracted["_market_context"] = None
                    extracted["_market_audit"] = {
                        "warnings": ["Market audit unavailable because postcode was not extracted."],
                    }
                    extracted["_pipeline_projects"] = _pipeline_projects
                    extracted["_pipeline_audit"] = _pipeline_audit
            except Exception as _de:
                print(f"[demand_service] error (non-fatal): {_de}")
                extracted["_demand_context"] = None
                extracted["_market_context"] = None
                extracted["_market_audit"] = {
                    "warnings": [f"Market audit unavailable because demand model failed: {_de}"],
                }
                try:
                    _pipeline_supply = build_pipeline_supply(req.pipelineProjects, req.pipelineIntel)
                    extracted["_pipeline_projects"] = _pipeline_supply["pipeline_projects"]
                    extracted["_pipeline_audit"] = _pipeline_supply["pipeline_audit"]
                except Exception:
                    extracted["_pipeline_projects"] = []
                    extracted["_pipeline_audit"] = None

            # ── Step 3: Score ─────────────────────────────────────────────
            yield sse_event("progress", {
                "step": 3, "total": 5,
                "label": "Scoring 17 dimensions",
                "detail": f"Analysing {centre_name}"
            })

            # Build pipeline intel context block if provided
            pipeline_intel_context = ""
            pi = req.pipelineIntel
            if pi:
                pipeline_intel_context = "\n\nPIPELINE INTEL (user-provided):\n"
                if pi.approved_das is not None:
                    pipeline_intel_context += f"- Approved DAs within 3km: {pi.approved_das}\n"
                if pi.lodged_applications is not None:
                    pipeline_intel_context += f"- Lodged applications: {pi.lodged_applications}\n"
                if pi.permit_sites is not None:
                    pipeline_intel_context += f"- Permit sites for sale: {pi.permit_sites}\n"
                if pi.notes:
                    pipeline_intel_context += f"- Notes: {pi.notes}\n"
                # Also attach extracted pipeline_mentions if any
                pipeline_mentions = extracted.get("pipeline_mentions", [])
                if pipeline_mentions:
                    pipeline_intel_context += f"- Document mentions: {'; '.join(pipeline_mentions)}\n"
                # Store a flag in scored data so the front-end knows pipeline intel was used
                extracted["_pipeline_intel_used"] = True
            if extracted.get("_pipeline_projects"):
                pipeline_intel_context += "\n\nSTRUCTURED DA / PIPELINE PROJECTS (manual):\n"
                for project in (extracted.get("_pipeline_projects") or [])[:10]:
                    pipeline_intel_context += (
                        f"- {project.get('status', 'unknown')}: "
                        f"{project.get('name') or project.get('address') or 'Pipeline project'}"
                        f" · {project.get('proposed_places') or 'unknown'} places"
                        f" · confidence {project.get('confidence', 'low')}\n"
                    )

            scoring_response = client.messages.create(
                model=MODEL,
                max_tokens=MAX_TOKENS,
                temperature=0,
                system=SCORING_SYSTEM_PROMPT,
                messages=[{
                    "role": "user",
                    "content": (
                        "Score this childcare centre acquisition.\n\n"
                        "IMPORTANT: Use normalised_ebitda for profitability and valuation scoring "
                        "if addbacks are present and confidence is 'high' or 'medium'. "
                        "Always state in the dimension summary whether you used reported or normalised EBITDA "
                        "and why.\n\n"
                        f"EXTRACTED DATA:\n{json.dumps(extracted, indent=2)}"
                        f"{pipeline_intel_context}"
                    )
                }]
            )
            scored_text = clean_json(scoring_response.content[0].text)
            scored      = json.loads(scored_text)

            # ── Server recalculates total_score — never trust Claude's value ──
            WEIGHTS = {
                'occupancy_demand':        0.15,
                'profitability_cashflow':  0.15,
                'revenue_pricing':         0.08,
                'staffing_resilience':     0.08,
                'lease_economics':         0.08,
                'valuation_structure':     0.08,
                'market_position':         0.07,
                'management_systems':      0.04,  # was 0.06; reduced to fund ccs_risk increase
                'regulatory_quality':      0.05,
                'upside_levers':           0.03,  # was 0.05; reduced to fund ccs_risk increase
                'ccs_risk':                0.07,  # was 0.03; increased — CCS risk is systemic
                'lease_tail':              0.03,
                'capex_liability':         0.02,
                'staff_qualification_mix': 0.02,
                'fee_benchmarking':        0.02,
                'operator_quality':        0.02,
                'enrolment_trend':         0.01,
                # Total: 1.00 ✓
            }
            dims = scored.get('dimensions', {})
            weighted_sum  = 0.0
            weight_used   = 0.0
            for dim_id, weight in WEIGHTS.items():
                dim = dims.get(dim_id, {})
                raw = dim.get('score')
                if isinstance(raw, (int, float)) and 0 <= raw <= 10:
                    weighted_sum += raw * weight
                    weight_used  += weight
            if weight_used > 0:
                # weighted_sum is in 0-10 range; scale to 0-100
                scored['total_score'] = round((weighted_sum / weight_used) * 10, 1)
            scored['scoring_version']  = '2.2'
            scored['scoring_timestamp'] = datetime.now(timezone.utc).isoformat()
            if req.pipelineIntel:
                scored['pipeline_intel_used'] = True
            if req.pipelineProjects:
                scored['pipeline_projects_used'] = True
            scored['pipeline_projects'] = extracted.get('_pipeline_projects') or []
            scored['pipeline_audit'] = extracted.get('_pipeline_audit')
            # Attach deterministic demand context to scored output
            # Frontend can surface EDR + zone alongside the score
            if extracted.get('_demand_context'):
                scored['effective_demand_ratio'] = extracted['_demand_context']['adj_kids_per_place']['mid']
                scored['demand_zone'] = extracted['_demand_context']['zone']
                scored['demand_context'] = extracted['_demand_context']
                scored['market_context'] = extracted.get('_market_context')
                scored['market_audit'] = extracted.get('_market_audit')

            structured_intel = build_structured_deal_intelligence(
                extracted=extracted,
                scored=scored,
                combined_text=combined_text,
                source_files=source_files,
                file_classes=file_classes,
            )

            # ── Step 4: Clean up ALL uploaded paths ───────────────────────
            yield sse_event("progress", {
                "step": 4, "total": 5,
                "label": "Generating report",
                "detail": "Mapping competitors · Building analysis"
            })

            # DEBUG/SAFETY: Do not delete uploads during the request.
            # Deleting here can cause HTTP 404 if the frontend retries, reconnects,
            # or double-submits the same storagePath before it receives the complete event.
            # Move cleanup to a scheduled job once the pipeline is stable.
            # try:
            #     supabase.storage.from_('uploads').remove(storage_paths)
            # except Exception:
            #     pass

            # ── Step 5: Complete ──────────────────────────────────────────
            yield sse_event("progress", {
                "step": 5, "total": 5,
                "label": "Complete",
                "detail": "Analysis ready"
            })

            yield sse_event("complete", {
                "success":   True,
                "extracted": extracted,
                "scored":    scored,
                "deal_summary": structured_intel["deal_summary"],
                "extracted_facts": structured_intel["extracted_facts"],
                "missing_fields": structured_intel["missing_fields"],
                "risks": structured_intel["risks"],
                "valuation_gate": structured_intel["valuation_gate"],
                "diligence_requests": structured_intel["diligence_requests"],
                "evidence": structured_intel["evidence"],
                "workflow": structured_intel,
                "retained_source_files": retained_source_files,
                "retention_warnings": retention_warnings,
                "meta": {
                    "source_files":  source_files,
                    "file_classes":  file_classes,
                    "skipped_files": skipped,
                    "retention_warnings": retention_warnings,
                    "pipeline_request_id": pipeline_request_id,
                    "retained_source_files": retained_source_files,
                }
            })

        except Exception as e:
            print(f"Pipeline error: {e}")
            yield sse_event("error", {"message": str(e)})
        finally:
            shutil.rmtree(work_dir, ignore_errors=True)

    return StreamingResponse(
        generate(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "X-Accel-Buffering": "no",
        }
    )
